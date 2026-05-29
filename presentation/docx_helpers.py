from __future__ import annotations

from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HE_FONT = "Arial"
CODE_FONT = "Consolas"

ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
ACCENT_DARK = RGBColor(0x0E, 0x4A, 0xA8)
DARK = RGBColor(0x1F, 0x23, 0x28)
MUTED = RGBColor(0x57, 0x60, 0x6A)
CODE_BG = RGBColor(0x27, 0x28, 0x22)
CODE_FG = RGBColor(0xF8, 0xF8, 0xF2)
CALLOUT_BG = RGBColor(0xEA, 0xF2, 0xFD)
TIP_BG = RGBColor(0xFF, 0xF8, 0xE1)
TIP_BORDER = RGBColor(0xF5, 0xA6, 0x23)
TABLE_HEAD_BG = RGBColor(0xEA, 0xF2, 0xFD)
WARN_BG = RGBColor(0xFF, 0xF8, 0xE1)
WARN_BORDER = RGBColor(0xF5, 0xA6, 0x23)


_PPR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
    "pPrChange",
)

_SECT_PR_ORDER = (
    "headerReference", "footerReference", "footnotePr", "endnotePr", "type",
    "pgSz", "pgMar", "paperSrc", "pgBorders", "lnNumType", "pgNumType",
    "cols", "formProt", "vAlign", "noEndnote", "titlePg", "textDirection",
    "bidi", "rtlGutter", "docGrid", "printerSettings", "sectPrChange",
)

_RPR_ORDER = (
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
    "dstrike", "outline", "shadow", "emboss", "imprint", "noProof",
    "snapToGrid", "vanish", "webHidden", "color", "spacing", "w", "kern",
    "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd",
    "fitText", "vertAlign", "rtl", "cs", "em", "lang", "eastAsianLayout",
    "specVanish", "oMath",
)

_TOGGLE_TAGS = frozenset({"bidi", "rtl", "rtlGutter"})


def _insert_ordered(parent, tag_name: str, order: tuple[str, ...]):
    existing = parent.find(qn(f"w:{tag_name}"))
    if existing is not None:
        if tag_name in _TOGGLE_TAGS:
            existing.set(qn("w:val"), "1")
        return existing
    target_idx = order.index(tag_name)
    new_el = OxmlElement(f"w:{tag_name}")
    if tag_name in _TOGGLE_TAGS:
        new_el.set(qn("w:val"), "1")
    for i, child in enumerate(list(parent)):
        local = child.tag.rsplit("}", 1)[-1]
        if local in order and order.index(local) > target_idx:
            parent.insert(i, new_el)
            return new_el
    parent.append(new_el)
    return new_el


def _insert_in_pPr_ordered(pPr, tag_name: str):
    return _insert_ordered(pPr, tag_name, _PPR_ORDER)


def _set_ltr(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = _insert_in_pPr_ordered(pPr, "bidi")
    bidi.set(qn("w:val"), "0")


def _set_bidi(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    _insert_in_pPr_ordered(pPr, "bidi")


def _set_run_rtl(run) -> None:
    rPr = run._r.get_or_add_rPr()
    rtl = _insert_ordered(rPr, "rtl", _RPR_ORDER)
    rtl.set(qn("w:val"), "1")


def _set_run_shading(run, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_paragraph_shading(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_paragraph_border(paragraph, side: str, color: str, sz: int = 12) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)
    pBdr.append(border)


def rgb_hex(rgb: RGBColor) -> str:
    return "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])


def _setup_section_rtl(section) -> None:
    sectPr = section._sectPr
    _insert_ordered(sectPr, "bidi", _SECT_PR_ORDER)
    _insert_ordered(sectPr, "rtlGutter", _SECT_PR_ORDER)


def _style_set_rtl(style, *, align_right: bool = True) -> None:
    style_el = style.element
    pPr = style_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_el.insert(0, pPr)
    _insert_in_pPr_ordered(pPr, "bidi")
    if align_right:
        jc = _insert_in_pPr_ordered(pPr, "jc")
        jc.set(qn("w:val"), "right")

    rPr = style_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_el.append(rPr)
    rtl = _insert_ordered(rPr, "rtl", _RPR_ORDER)
    rtl.set(qn("w:val"), "1")


def _set_doc_defaults_rtl(doc: Document) -> None:
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        return

    pPrDefault = docDefaults.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        pPrDefault = OxmlElement("w:pPrDefault")
        docDefaults.append(pPrDefault)
    pPr = pPrDefault.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDefault.append(pPr)
    _insert_in_pPr_ordered(pPr, "bidi")
    jc = _insert_in_pPr_ordered(pPr, "jc")
    jc.set(qn("w:val"), "right")

    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)
    rtl = _insert_ordered(rPr, "rtl", _RPR_ORDER)
    rtl.set(qn("w:val"), "1")
    lang = _insert_ordered(rPr, "lang", _RPR_ORDER)
    lang.set(qn("w:bidi"), "he-IL")


def setup_styles(doc: Document) -> None:
    _set_doc_defaults_rtl(doc)

    normal = doc.styles["Normal"]
    normal.font.name = HE_FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = DARK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), HE_FONT)
    rfonts.set(qn("w:ascii"), HE_FONT)
    rfonts.set(qn("w:hAnsi"), HE_FONT)
    _style_set_rtl(normal, align_right=True)

    for heading_name, size, color in [
        ("Heading 1", 22, ACCENT_DARK),
        ("Heading 2", 17, ACCENT),
        ("Heading 3", 14, DARK),
    ]:
        s = doc.styles[heading_name]
        s.font.name = HE_FONT
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        _style_set_rtl(s, align_right=True)

    for list_style in ("List Bullet", "List Number"):
        try:
            s = doc.styles[list_style]
        except KeyError:
            continue
        s.font.name = HE_FONT
        _style_set_rtl(s, align_right=True)


def he_paragraph(doc, text=None, *, style=None, bold=False, size=12,
                 italic=False, color=None, space_before=2, space_after=4,
                 align=WD_ALIGN_PARAGRAPH.RIGHT):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    _set_bidi(p)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text is None:
        return p
    run = p.add_run(text)
    _set_run_rtl(run)
    run.font.name = HE_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def he_segments(doc, segments, *, style=None, size=12, bullet=False,
                bold=False, space_before=2, space_after=4,
                align=WD_ALIGN_PARAGRAPH.RIGHT):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    elif style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    _set_bidi(p)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    for seg in segments:
        text, kind = seg if isinstance(seg, tuple) else (seg, "he")
        run = p.add_run(text)
        if kind == "en":
            run.font.name = CODE_FONT
            run.font.size = Pt(size - 1)
            run.font.color.rgb = RGBColor(0x03, 0x2F, 0x62)
            _set_run_shading(run, "F1F3F5")
        elif kind == "bold":
            run.font.name = HE_FONT
            run.font.size = Pt(size)
            run.font.bold = True
            _set_run_rtl(run)
        elif kind == "italic":
            run.font.name = HE_FONT
            run.font.size = Pt(size)
            run.font.italic = True
            _set_run_rtl(run)
        elif kind == "accent":
            run.font.name = HE_FONT
            run.font.size = Pt(size)
            run.font.bold = True
            run.font.color.rgb = ACCENT
            _set_run_rtl(run)
        elif kind == "muted":
            run.font.name = HE_FONT
            run.font.size = Pt(size)
            run.font.color.rgb = MUTED
            _set_run_rtl(run)
        else:
            run.font.name = HE_FONT
            run.font.size = Pt(size)
            run.font.bold = bold
            _set_run_rtl(run)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    _set_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _set_run_rtl(run)
    run.font.name = HE_FONT
    run.font.size = Pt(22 if level == 1 else 17 if level == 2 else 14)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK if level == 1 else (ACCENT if level == 2 else DARK)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def callout(doc, segments, *, bg=CALLOUT_BG, border_color=ACCENT, size=11):
    p = he_segments(doc, segments, size=size, space_before=4, space_after=4)
    _set_paragraph_shading(p, rgb_hex(bg))
    _add_paragraph_border(p, "right", rgb_hex(border_color), sz=18)
    _add_paragraph_border(p, "top", rgb_hex(border_color), sz=4)
    _add_paragraph_border(p, "bottom", rgb_hex(border_color), sz=4)
    _add_paragraph_border(p, "left", rgb_hex(border_color), sz=4)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    return p


def tip_box(doc, segments):
    return callout(doc, segments, bg=TIP_BG, border_color=TIP_BORDER, size=11)


def warn_box(doc, text: str):
    p = he_paragraph(doc, text, size=11, bold=True, color=DARK,
                     space_before=4, space_after=4)
    _set_paragraph_shading(p, rgb_hex(WARN_BG))
    _add_paragraph_border(p, "right", rgb_hex(WARN_BORDER), sz=18)
    _add_paragraph_border(p, "top", rgb_hex(WARN_BORDER), sz=4)
    _add_paragraph_border(p, "bottom", rgb_hex(WARN_BORDER), sz=4)
    _add_paragraph_border(p, "left", rgb_hex(WARN_BORDER), sz=4)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    return p


def code_block(doc, lines: Sequence[str], *, size: int = 9):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    cell = table.cell(0, 0)
    _set_cell_shading(cell, rgb_hex(CODE_BG))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.text = ""

    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _set_ltr(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(size)
        run.font.color.rgb = CODE_FG
    return table


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]],
              col_widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = True

    bidi = OxmlElement("w:bidiVisual")
    table._tbl.tblPr.append(bidi)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, rgb_hex(TABLE_HEAD_BG))
        cell.text = ""
        p = cell.paragraphs[0]
        _set_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(header)
        _set_run_rtl(run)
        run.font.name = HE_FONT
        run.font.size = Pt(11)
        run.font.bold = True

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _set_bidi(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(value))
            _set_run_rtl(run)
            run.font.name = HE_FONT
            run.font.size = Pt(10)

    if col_widths:
        for col, w in zip(table.columns, col_widths):
            for cell in col.cells:
                cell.width = Cm(w)


def he_bullet(doc, text: str, *, size=12):
    return he_paragraph(doc, "\u25CF  " + text, size=size,
                        space_before=1, space_after=1)


def new_document() -> Document:
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    _setup_section_rtl(section)
    return doc


def add_page_numbers(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = HE_FONT
    run.font.size = Pt(10)


def word_postprocess_rtl(docx_path: Path) -> None:
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        return

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdocx = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False)
        try:
            try:
                for toc in wdocx.TablesOfContents:
                    toc.Update()
            except Exception:
                pass

            for para in wdocx.Paragraphs:
                try:
                    current_align = para.Format.Alignment
                except Exception:
                    current_align = None
                if current_align == 1:
                    continue
                font_name = ""
                try:
                    font_name = (para.Range.Font.Name or "") + " " + \
                                (para.Range.Font.NameAscii or "")
                except Exception:
                    font_name = ""
                if "Consolas" in font_name or "Courier" in font_name:
                    continue
                try:
                    para.Format.ReadingOrder = 1
                except Exception:
                    pass
                if current_align in (0, None):
                    try:
                        para.Format.Alignment = 2
                    except Exception:
                        pass
            wdocx.Save()
        finally:
            wdocx.Close(SaveChanges=False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
