from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_helpers import (
    ACCENT, ACCENT_DARK, DARK, MUTED,
    add_page_numbers, add_table, callout, code_block, he_bullet,
    he_paragraph, he_segments, heading, new_document, page_break,
    set_update_fields_on_open, tip_box, word_postprocess_rtl, _set_bidi,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = REPO_ROOT / "מדריך-הסבר-הפרויקט-לתלמיד.docx"


def build_document() -> None:
    doc = new_document()
    add_page_numbers(doc)

    write_cover(doc)
    page_break(doc)

    write_toc(doc)
    page_break(doc)

    heading(doc, "0. איך ללמוד מהמדריך הזה", level=1)
    write_how_to_use(doc)
    page_break(doc)

    heading(doc, "1. מבט-על: איך כל החלקים מתחברים", level=1)
    write_overview(doc)
    page_break(doc)

    heading(doc, "2. נימות (Threads) — ההסבר המלא", level=1)
    write_threads(doc)
    page_break(doc)

    heading(doc, "3. תקשורת רשת: סוקטים ו-TCP", level=1)
    write_network(doc)
    page_break(doc)

    heading(doc, "4. ערבול סיסמאות (SHA-256 + salt)", level=1)
    write_hashing(doc)
    page_break(doc)

    heading(doc, "5. תכנות מונחה עצמים (OOP)", level=1)
    write_oop(doc)
    page_break(doc)

    heading(doc, "6. פרוטוקול ההודעות ו-JSON", level=1)
    write_protocol(doc)
    page_break(doc)

    heading(doc, "7. מבנה השרת: מי אחראי על מה", level=1)
    write_server_structure(doc)
    page_break(doc)

    heading(doc, "8. הבינה המלאכותית (ChessAI)", level=1)
    write_ai(doc)
    page_break(doc)

    heading(doc, "9. חוקי השחמט שמומשו (Board)", level=1)
    write_chess_rules(doc)
    page_break(doc)

    heading(doc, "10. ממשק המשתמש הגרפי (GUI / Swing)", level=1)
    write_gui(doc)
    page_break(doc)

    heading(doc, "11. מילון מונחים להגנה", level=1)
    write_glossary(doc)
    page_break(doc)

    heading(doc, "12. בנק שאלות ותשובות מסכם", level=1)
    write_qa_bank(doc)

    set_update_fields_on_open(doc)
    doc.save(str(OUT_DOCX))
    word_postprocess_rtl(OUT_DOCX)
    print(f"wrote study guide docx ({OUT_DOCX.stat().st_size:,} bytes)")


def write_cover(doc):
    he_paragraph(doc, "", space_before=60)
    he_paragraph(doc, "מדריך הסבר הפרויקט — לתלמיד", size=26, bold=True, color=ACCENT_DARK,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    he_paragraph(doc, "משחק שחמט מאובטח בין שני משתתפים (Java)", size=20, bold=True, color=ACCENT,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    he_paragraph(doc,
                 "מדריך לימוד שמסביר — בשפה פשוטה ועם דוגמאות קוד אמיתיות מהפרויקט — "
                 "איך כל רכיב עובד ואיך לספר עליו בהגנה בעל-פה.",
                 size=13, color=MUTED, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    he_paragraph(doc, "להבנה ולשינון לפני הצגת הפרויקט", size=12, color=DARK,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)


def write_toc(doc):
    heading(doc, "תוכן עניינים", level=1)
    he_paragraph(doc,
                 "תוכן העניינים מתעדכן אוטומטית בפתיחת הקובץ ב-Word. "
                 "לעדכון ידני: קליק-ימני על הטבלה ובחירת 'Update Field' (או F9).",
                 italic=True, color=MUTED)

    p = doc.add_paragraph()
    _set_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u \\p " "'
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r = p.add_run()._r
    r.append(fld1)
    r.append(instr)
    r.append(fld2)
    r.append(fld3)
    he_paragraph(doc, "תוכן עניינים — לחץ F9 ב-Word כדי לעדכן", italic=True, color=MUTED)


def write_how_to_use(doc):
    he_paragraph(doc,
                 "המדריך בנוי כך שכל פרק מסביר נושא אחד בשלושה שלבים: קודם הרעיון "
                 "בשפה פשוטה (לפעמים עם אנלוגיה מהחיים), אחר כך הקוד האמיתי מהפרויקט "
                 "שמראה איפה זה קורה, ובסוף 'איך להגיד את זה בהגנה' — משפטים מוכנים "
                 "שאפשר לומר לבוחן.")
    he_bullet(doc, "קרא קודם את ההסבר הפשוט עד שאתה מבין את הרעיון.")
    he_bullet(doc, "אחר כך הצמד את הקוד לרעיון — איזו שורה עושה מה.")
    he_bullet(doc, "תרגל בקול רם את החלק 'איך להגיד בהגנה'.")
    he_bullet(doc, "בסוף כל פרק יש שאלות ותשובות — כסה את התשובה ונסה לענות לבד.")
    tip_box(doc, [
        ("טיפ זהב: ", "bold"),
        ("הבוחן לא מצפה שתדע הכול בעל-פה מילה במילה. הוא רוצה לראות שאתה "
         "מבין למה בחרת כל דבר ואיך הוא עובד. תמיד תסביר 'מה הבעיה שזה פותר' "
         "לפני 'איך מימשתי'.", "he"),
    ])


def write_overview(doc):
    he_paragraph(doc,
                 "הפרויקט הוא משחק שחמט שמשוחק בין שני מחשבים דרך הרשת. יש תוכנית "
                 "אחת שנקראת 'שרת' (Server) שיושבת במרכז ומנהלת את הכול, ויש תוכניות "
                 "'לקוח' (Client) — אחת לכל שחקן — שמציגות את הלוח ושולחות מהלכים.")
    he_paragraph(doc, "זרימת המשחק בקצרה:")
    he_bullet(doc, "השרת עולה ומחכה לחיבורים בכתובת 127.0.0.1 ובפורט 5555.")
    he_bullet(doc, "כל לקוח מתחבר לשרת, נרשם או מתחבר עם שם משתמש וסיסמה.")
    he_bullet(doc, "השחקן לוחץ 'שחק נגד אדם' (ונכנס לתור) או 'שחק נגד מחשב'.")
    he_bullet(doc, "כששני שחקנים בתור — נפתח משחק, והשרת שולח לשניהם את מצב הלוח.")
    he_bullet(doc, "כל מהלך נשלח לשרת, השרת בודק שהוא חוקי, מעדכן את הלוח, ומחזיר את המצב החדש לשני הצדדים.")
    callout(doc, [
        ("המשפט שמסכם הכול: ", "bold"),
        ("\"השרת הוא הבוס — הוא מחזיק את הלוח האמיתי, בודק כל מהלך, ומודיע "
         "לשני הלקוחות מה קרה. הלקוח רק מצייר ושולח בקשות.\"", "he"),
    ])
    he_paragraph(doc, "החלוקה לקבצים העיקריים:")
    add_table(doc, ["קובץ", "תפקיד בקצרה"], [
        ("Main.java", "נקודת הכניסה — מריץ שרת או לקוח"),
        ("server/ChessServer.java", "מאזין לחיבורים ופותח נימה לכל לקוח"),
        ("server/ClientHandler.java", "מטפל בלקוח אחד — קורא הודעות ומגיב"),
        ("server/Matchmaker.java", "תור שמזווג שני שחקנים"),
        ("server/Game.java", "משחק בודד — מחזיק לוח ושני שחקנים"),
        ("Board.java", "הלוח, חוקי התנועה, שח/מט/פט"),
        ("ChessAI.java", "בחירת מהלך למחשב"),
        ("UserAuth.java", "הרשמה/התחברות + ערבול סיסמאות"),
        ("JSONUtil.java", "קידוד ופענוח הודעות JSON"),
        ("client/NetworkClient.java", "צד הלקוח של החיבור לשרת"),
        ("client/ChessGUI.java", "המסכים והלוח שהשחקן רואה (Swing)"),
    ])


def write_threads(doc):
    heading(doc, "2.1 הבעיה: למה בכלל צריך נימות?", level=2)
    he_paragraph(doc,
                 "תאר לך מלצר אחד במסעדה שחייב לסיים לטפל בשולחן אחד עד הסוף לפני "
                 "שהוא ניגש לשולחן הבא. אם לקוח אחד מתלבט חצי שעה — כל השאר רעבים "
                 "ומחכים. זו בדיוק הבעיה בשרת: אם השרת יטפל בלקוח אחד בלבד בכל רגע, "
                 "שחקן שחושב על המהלך שלו יתקע את כל שאר השחקנים.")
    he_paragraph(doc,
                 "הפתרון: נימה (Thread) נפרדת לכל לקוח. זה כמו להעסיק מלצר נפרד לכל "
                 "שולחן — כולם מטופלים במקביל, ולקוח איטי לא חוסם אף אחד.")

    heading(doc, "2.2 מה זה נימה (Thread) בעצם?", level=2)
    he_paragraph(doc, "כדי להבין נימה, צריך קודם להבין את ההבדל מתהליך:")
    he_bullet(doc, "תהליך (Process) — תוכנית רצה עם זיכרון נפרד משלה (למשל כל חלון Word).")
    he_bullet(doc, "נימה (Thread) — 'נתיב ריצה' בתוך תוכנית אחת. כל הנימות באותה תוכנית חולקות את אותו זיכרון, ולכן יכולות לגשת לאותם נתונים.")
    he_paragraph(doc,
                 "ב-Java נימה מריצה קוד שנמצא במתודה run(). יוצרים אובייקט שמממש את "
                 "הממשק Runnable, עוטפים אותו ב-Thread, וקוראים start(). מאותו רגע "
                 "הקוד שב-run() רץ 'במקביל' לשאר התוכנית.")

    heading(doc, "2.3 איך זה מומש בפרויקט — נימה לכל לקוח", level=2)
    he_paragraph(doc,
                 "בלב השרת יש 'לולאת קבלה' (accept loop). השרת קורא ל-accept() "
                 "שמחזיר חיבור (Socket) ברגע שלקוח חדש מתחבר. מיד אחרי זה השרת פותח "
                 "נימה חדשה שתטפל בלקוח הזה, וחוזר מיד להמתין ללקוח הבא:")
    code_block(doc, [
        "// ChessServer.java",
        "try (ServerSocket server = new ServerSocket(PORT)) {   // PORT = 5555",
        "    while (true) {",
        "        Socket socket = server.accept();               // ממתין ללקוח חדש",
        "        ClientHandler handler = new ClientHandler(socket, auth, matchmaker);",
        "        new Thread(handler).start();                   // נימה נפרדת ללקוח",
        "    }",
        "}",
    ])
    he_paragraph(doc,
                 "המחלקה ClientHandler מממשת Runnable, ולכן אפשר להריץ אותה בנימה. "
                 "הקוד שרץ בנימה נמצא ב-run(): לולאה שקוראת שורה-שורה מהלקוח, "
                 "ומפעילה את הטיפול המתאים לכל הודעה:")
    code_block(doc, [
        "// ClientHandler.java — implements Runnable",
        "public void run() {",
        "    try {",
        "        String line;",
        "        while ((line = in.readLine()) != null) {   // קריאה חוסמת עד שמגיעה שורה",
        "            handleMessage(JSONUtil.decode(line));   // טיפול בהודעה",
        "        }",
        "    } catch (IOException ignored) {",
        "    } finally {",
        "        cleanup();   // הסרה מהתור וסגירת החיבור כשהלקוח מתנתק",
        "    }",
        "}",
    ])
    callout(doc, [
        ("שים לב: ", "bold"),
        ("readLine() היא פעולה 'חוסמת' — הנימה פשוט מחכה שם בלי לעשות כלום עד "
         "שמגיעה שורה. בזכות שזו נימה נפרדת, ההמתנה הזו לא מפריעה לאף לקוח אחר.", "he"),
    ])

    heading(doc, "2.4 הבעיה הנסתרת: מירוץ על נתונים משותפים", level=2)
    he_paragraph(doc,
                 "מכיוון שכל הנימות חולקות זיכרון, עלולה להיווצר בעיה: מה אם שתי "
                 "נימות (שני השחקנים) ינסו לשנות את אותו לוח באותו רגע? המהלכים "
                 "עלולים 'להתנגש' והנתונים יתקלקלו. זה נקרא 'תנאי מירוץ' (race condition).")
    he_paragraph(doc,
                 "הפתרון ב-Java הוא המילה synchronized. מתודה שמסומנת synchronized "
                 "מאפשרת רק לנימה אחת להיכנס אליה בכל רגע — השאר ממתינות בתור. כך "
                 "מובטח שמהלך אחד מסתיים לפני שמתחיל הבא:")
    code_block(doc, [
        "// Game.java — רק נימה אחת מבצעת מהלך בכל רגע",
        "public synchronized boolean applyMove(ClientHandler handler,",
        "                                      int fc, int fr, int tc, int tr) {",
        "    Color playerColor = (handler == white) ? Color.WHITE : Color.BLACK;",
        "    if (board.turn != playerColor) return false;   // לא תורך",
        "    if (!board.isLegalMove(fc, fr, tc, tr)) return false;",
        "    board.applyMove(fc, fr, tc, tr);",
        "    afterMove();",
        "    return true;",
        "}",
    ])
    he_paragraph(doc,
                 "אותו רעיון נמצא גם ב-Matchmaker (התור), שבו שתי נימות עלולות לנסות "
                 "להיכנס לתור באותו רגע, וגם ב-UserAuth (קריאה/כתיבה לקובץ המשתמשים):")
    code_block(doc, [
        "// Matchmaker.java — זיווג בטוח של שני שחקנים",
        "public synchronized ClientHandler tryMatch(ClientHandler client) {",
        "    if (!queue.isEmpty()) {",
        "        return queue.poll();      // יש יריב ממתין — מחזירים אותו",
        "    }",
        "    queue.add(client);            // אין יריב — נכנסים לתור",
        "    return null;",
        "}",
    ])

    heading(doc, "2.5 נימה גם בצד הלקוח", level=2)
    he_paragraph(doc,
                 "גם הלקוח משתמש בנימה: כדי שהמסך לא 'יקפא' בזמן שהוא ממתין להודעות "
                 "מהשרת, הלקוח מריץ נימת-רקע (daemon) שכל תפקידה לקרוא הודעות נכנסות "
                 "ולשים אותן בתור (inbox). המסך הראשי ממשיך להגיב למשתמש:")
    code_block(doc, [
        "// NetworkClient.java",
        "Thread reader = new Thread(() -> {",
        "    String line;",
        "    while ((line = in.readLine()) != null) {",
        "        inbox.put(JSONUtil.decode(line));   // BlockingQueue בטוח לנימות",
        "    }",
        "});",
        "reader.setDaemon(true);   // נימת רקע — נסגרת אוטומטית עם התוכנית",
        "reader.start();",
    ])

    heading(doc, "2.6 איך להגיד את זה בהגנה", level=2)
    he_segments(doc, [("משפט פתיחה: ", "accent"),
                      ("\"השתמשתי בנימה נפרדת לכל לקוח כדי שהשרת יטפל בכמה שחקנים "
                       "במקביל, ולקוח איטי לא יחסום את השאר.\"", "he")], size=12)
    he_segments(doc, [("איך יצרתי: ", "accent"),
                      ("\"ClientHandler מממש Runnable, ובלולאת ה-accept אני עושה "
                       "new Thread(handler).start().\"", "he")], size=12)
    he_segments(doc, [("הגנה על מידע: ", "accent"),
                      ("\"כי כל הנימות חולקות זיכרון, סימנתי את מתודות המשחק והתור "
                       "כ-synchronized, כך שרק נימה אחת נכנסת בכל פעם והנתונים לא "
                       "מתקלקלים.\"", "he")], size=12, space_after=8)

    heading(doc, "2.7 שאלות ותשובות — נימות", level=2)
    qa_list(doc, [
        ("מה ההבדל בין תהליך לנימה?",
         "תהליך הוא תוכנית עם זיכרון נפרד. נימה היא נתיב ריצה בתוך תוכנית, וכל "
         "הנימות חולקות את אותו זיכרון."),
        ("איך פותחים נימה ב-Java?",
         "יוצרים אובייקט שמממש Runnable (אצלי ClientHandler), עוטפים ב-new Thread(...) "
         "וקוראים start(). הקוד רץ במתודה run()."),
        ("למה צריך synchronized?",
         "כי שתי נימות יכולות לגשת לאותו מידע משותף בו-זמנית ולקלקל אותו. synchronized "
         "מבטיח שרק נימה אחת נכנסת למתודה בכל רגע."),
        ("מה קורה אם לקוח מתנתק?",
         "readLine() מחזיר null או נזרקת חריגה, הלולאה נגמרת, וב-finally קוראים "
         "ל-cleanup() שמסיר אותו מהתור וסוגר את החיבור — בלי להפיל את השרת."),
    ])


def write_network(doc):
    heading(doc, "3.1 הרעיון: מחשבים שמדברים", level=2)
    he_paragraph(doc,
                 "כדי ששני מחשבים ידברו, צריך 'כתובת' ו'ערוץ'. כתובת ה-IP מזהה את "
                 "המחשב ברשת (בפרויקט 127.0.0.1 = 'המחשב הזה'), והפורט (5555) מזהה "
                 "איזו תוכנית בתוך המחשב צריכה לקבל את ההודעה. יחד הם אומרים בדיוק "
                 "לאן לשלוח.")
    he_paragraph(doc,
                 "סוקט (Socket) הוא 'קצה הצינור' — אובייקט שדרכו שולחים ומקבלים "
                 "נתונים. בצד השרת יש ServerSocket שמאזין, ובצד הלקוח יש Socket "
                 "שמתחבר.")

    heading(doc, "3.2 למה TCP ולא UDP?", level=2)
    he_bullet(doc, "TCP — אמין ומסודר: כל הודעה מגיעה, ובסדר הנכון. איטי מעט יותר.")
    he_bullet(doc, "UDP — מהיר אך לא אמין: הודעות יכולות ללכת לאיבוד או להגיע בלי סדר.")
    he_paragraph(doc,
                 "בשחמט בחרתי TCP, כי אסור שמהלך 'יאבד' או יגיע אחרי מהלך מאוחר "
                 "יותר — זה יהרוס את המשחק.")

    heading(doc, "3.3 איך זה מומש — צד השרת", level=2)
    he_paragraph(doc, "השרת פותח ServerSocket ובלולאה קורא accept() לכל חיבור חדש:")
    code_block(doc, [
        "// ChessServer.java",
        "ServerSocket server = new ServerSocket(5555);",
        "Socket socket = server.accept();   // מחזיר חיבור ללקוח שהתחבר",
    ])
    he_paragraph(doc,
                 "כל ClientHandler עוטף את הסוקט בשני זרמי טקסט: PrintWriter לכתיבה "
                 "ו-BufferedReader לקריאה, שניהם בקידוד UTF-8 (כדי לתמוך בעברית):")
    code_block(doc, [
        "// ClientHandler.java",
        "out = new PrintWriter(new OutputStreamWriter(",
        "          socket.getOutputStream(), \"UTF-8\"), true);",
        "in  = new BufferedReader(new InputStreamReader(",
        "          socket.getInputStream(), \"UTF-8\"));",
    ])

    heading(doc, "3.4 איך זה מומש — צד הלקוח", level=2)
    he_paragraph(doc, "הלקוח יוצר Socket שמתחבר לכתובת ולפורט של השרת:")
    code_block(doc, [
        "// NetworkClient.java",
        "private static final String HOST = \"127.0.0.1\";",
        "private static final int PORT = 5555;",
        "...",
        "socket = new Socket(HOST, PORT);   // מבצע connect לשרת",
    ])
    callout(doc, [
        ("שורת מפתח אחת לזכור: ", "bold"),
        ("בצד השרת — server.accept(). בצד הלקוח — new Socket(host, port). "
         "אלה שתי הפעולות שמקימות את החיבור.", "he"),
    ])

    heading(doc, "3.5 שאלות ותשובות — תקשורת", level=2)
    qa_list(doc, [
        ("מה זה כתובת IP ומה זה פורט?",
         "IP מזהה מחשב ברשת. פורט מזהה תוכנית מסוימת באותו מחשב (אצלי 5555)."),
        ("מהם השלבים בפתיחת שרת ב-Java?",
         "יוצרים ServerSocket עם פורט, ובלולאה קוראים accept() שמחזיר Socket לכל "
         "לקוח שמתחבר."),
        ("מה עושה הלקוח כדי להתחבר?",
         "יוצר new Socket(host, port) — זה מבצע connect לפי כתובת ה-IP והפורט."),
        ("למה השתמשת ב-UTF-8?",
         "כדי שהודעות עם תווים בעברית (וכל תו מיוחד) יעברו נכון בין השרת ללקוח."),
    ])


def write_hashing(doc):
    heading(doc, "4.1 הרעיון: לא לשמור סיסמה גלויה", level=2)
    he_paragraph(doc,
                 "אם נשמור את הסיסמאות כפי שהן בקובץ, מי שיגנוב את הקובץ ידע את "
                 "כל הסיסמאות. לכן לא שומרים את הסיסמה עצמה, אלא 'טביעת אצבע' שלה — "
                 "תוצאה של פונקציית ערבול (hash).")
    he_paragraph(doc,
                 "פונקציית ערבול היא חד-כיוונית: קל לחשב ממנה תוצאה מתוך הסיסמה, "
                 "אבל כמעט בלתי אפשרי לחזור מהתוצאה אל הסיסמה. בפרויקט אני משתמש "
                 "ב-SHA-256, פונקציית ערבול סטנדרטית, דרך המחלקה MessageDigest.")

    heading(doc, "4.2 מה זה salt ולמה הוא קריטי", level=2)
    he_paragraph(doc,
                 "אם שני אנשים יבחרו אותה סיסמה, הערבול שלהם יֵצא זהה — וזה מסגיר "
                 "מידע. בנוסף, קיימות טבלאות מוכנות מראש (rainbow tables) שמתרגמות "
                 "ערבולים נפוצים בחזרה לסיסמאות. כדי לחסום את זה מוסיפים לכל משתמש "
                 "salt — מחרוזת אקראית ייחודית — לפני הערבול. כך גם סיסמאות זהות "
                 "מקבלות תוצאה שונה לגמרי.")

    heading(doc, "4.3 איך זה מומש — הרשמה", level=2)
    code_block(doc, [
        "// UserAuth.java — register()",
        "byte[] saltBytes = new byte[16];",
        "new SecureRandom().nextBytes(saltBytes);                 // salt אקראי",
        "String salt = Base64.getEncoder().encodeToString(saltBytes);",
        "String hash = hash(password, salt);                      // SHA-256(salt+password)",
        "users.put(username, new String[]{salt, hash});           // שומרים salt + hash בלבד",
        "save();                                                  // לקובץ users.json",
    ])
    he_paragraph(doc, "פונקציית הערבול עצמה:")
    code_block(doc, [
        "// UserAuth.java — hash()",
        "MessageDigest md = MessageDigest.getInstance(\"SHA-256\");",
        "String input = salt + password;",
        "byte[] digest = md.digest(input.getBytes(StandardCharsets.UTF_8));",
        "// ממירים את הבייטים למחרוזת הקסדצימלית ומחזירים",
    ])

    heading(doc, "4.4 איך זה מומש — התחברות", level=2)
    he_paragraph(doc,
                 "בהתחברות לא משווים סיסמאות. לוקחים את ה-salt השמור של המשתמש, "
                 "מחשבים מחדש את הערבול של הסיסמה שהוזנה, ומשווים לערבול השמור. "
                 "אם הם זהים — הסיסמה נכונה:")
    code_block(doc, [
        "// UserAuth.java — login()",
        "String[] entry = users.get(username);   // entry = {salt, hash}",
        "if (entry == null) return false;",
        "return entry[1].equals(hash(password, entry[0]));",
    ])
    callout(doc, [
        ("ההבחנה שחייבים לדעת: ", "bold"),
        ("ערבול הוא לא הצפנה. הצפנה היא דו-כיוונית (אפשר לפענח בחזרה עם מפתח), "
         "ערבול הוא חד-כיווני. לסיסמאות משתמשים בערבול, כי השרת לא צריך לדעת את "
         "הסיסמה המקורית — רק לבדוק אם מה שהוזן נותן את אותה תוצאה.", "he"),
    ])
    he_paragraph(doc,
                 "אפשר להוכיח זאת בקלות: פותחים את users.json ורואים שאין שם סיסמאות "
                 "גלויות — רק salt ו-hash מופרדים בנקודתיים.")

    heading(doc, "4.5 שאלות ותשובות — ערבול", level=2)
    qa_list(doc, [
        ("מה ההבדל בין ערבול להצפנה?",
         "הצפנה דו-כיוונית — אפשר לפענח בחזרה עם מפתח. ערבול חד-כיווני — אי אפשר "
         "לשחזר את הקלט. לסיסמאות משתמשים בערבול."),
        ("מה זה salt ולמה הוא חשוב?",
         "מחרוזת אקראית שמתווספת לסיסמה לפני הערבול. בלעדיו, שתי סיסמאות זהות "
         "מקבלות תוצאה זהה, וזה מאפשר התקפה בעזרת טבלאות מוכנות."),
        ("איך נשמרת הסיסמה בפרויקט?",
         "לא נשמרת הסיסמה עצמה — רק ה-salt והערבול של (salt + סיסמה) בקובץ "
         "users.json. בהתחברות מחשבים שוב ומשווים."),
        ("מה זה SecureRandom?",
         "מחולל מספרים אקראיים חזק מבחינה קריפטוגרפית, שמתאים ליצירת salt בטוח."),
    ])


def write_oop(doc):
    heading(doc, "5.1 ארבעת העקרונות — בקצרה", level=2)
    add_table(doc, ["עיקרון", "מה זה", "איפה בפרויקט"], [
        ("הפשטה (Abstraction)", "מחלקת אב שמגדירה 'מה' בלי 'איך'", "Piece (abstract)"),
        ("ירושה (Inheritance)", "מחלקה יורשת תכונות מאב", "Knight/Rook/... extends Piece"),
        ("פולימורפיזם", "אותה קריאה מתנהגת אחרת לכל סוג", "getMoves() לכל כלי"),
        ("כימוס (Encapsulation)", "הסתרת הפרטים הפנימיים", "grid פרטי ב-Board"),
    ])

    heading(doc, "5.2 הפשטה וירושה — מחלקת Piece", level=2)
    he_paragraph(doc,
                 "כל כלי שחמט הוא 'כלי', אבל לכל אחד חוקי תנועה משלו. לכן יצרתי "
                 "מחלקת אב מופשטת Piece שמגדירה את מה שמשותף לכולם, ומשאירה לכל כלי "
                 "לממש את התנועה שלו. abstract אומר: 'אי אפשר ליצור Piece סתם, חייבים "
                 "כלי קונקרטי':")
    code_block(doc, [
        "// Piece.java — מחלקת אב מופשטת",
        "public abstract class Piece {",
        "    public final Color color;",
        "    public boolean hasMoved;",
        "    public abstract List<int[]> getMoves(Board board, int col, int row);",
        "    public abstract int getValue();",
        "    public abstract String getSymbol();",
        "}",
    ])
    he_paragraph(doc,
                 "כל כלי יורש מ-Piece ומממש את getMoves בדרכו. הפרש (Knight), למשל, "
                 "'קופץ' בצורת L, ולכן הוא משתמש בעזרת jumpMoves עם רשימת קפיצות:")
    code_block(doc, [
        "// Knight.java — יורש מ-Piece",
        "public class Knight extends Piece {",
        "    private static final int[][] OFFSETS = {",
        "        {1,2},{2,1},{2,-1},{1,-2},{-1,-2},{-2,-1},{-2,1},{-1,2}",
        "    };",
        "    public Knight(Color color) { super(color); }",
        "    @Override public List<int[]> getMoves(Board b, int col, int row) {",
        "        return jumpMoves(b, col, row, OFFSETS);",
        "    }",
        "}",
    ])
    he_paragraph(doc,
                 "כלים שזזים 'בקו' (צריח, רץ, מלכה) חולקים לוגיקה אחרת — slideMoves — "
                 "שמוגדרת פעם אחת במחלקת האב Piece, וכך נמנע כפל קוד:")
    code_block(doc, [
        "// Piece.java — תנועה בקו עד שנתקלים בכלי או בקצה",
        "protected List<int[]> slideMoves(Board board, int col, int row, int[][] dirs) {",
        "    List<int[]> moves = new ArrayList<>();",
        "    for (int[] d : dirs) {",
        "        int step = 1;",
        "        while (true) {",
        "            int tc = col + d[0]*step, tr = row + d[1]*step;",
        "            if (!inBounds(tc, tr)) break;",
        "            Piece target = board.getPiece(tc, tr);",
        "            if (target == null) { moves.add(new int[]{tc, tr}); }",
        "            else { if (target.color != color) moves.add(new int[]{tc, tr}); break; }",
        "            step++;",
        "        }",
        "    }",
        "    return moves;",
        "}",
    ])

    heading(doc, "5.3 פולימורפיזם — הקסם", level=2)
    he_paragraph(doc,
                 "כשהלוח רוצה לדעת לאן כלי יכול לזוז, הוא פשוט קורא piece.getMoves(...) "
                 "— בלי לדעת אם זה פרש, צריח או מלכה. כל כלי 'יודע בעצמו' איך הוא זז. "
                 "זה פולימורפיזם: אותה שורת קוד מתנהגת נכון לכל סוג כלי. אם נוסיף "
                 "כלי חדש בעתיד, הלוח לא ישתנה בכלל.")

    heading(doc, "5.4 איך להגיד את זה בהגנה", level=2)
    he_segments(doc, [("\"", "he"),
                      ("יצרתי מחלקת אב מופשטת Piece עם getMoves מופשטת. כל כלי יורש "
                       "ממנה ומממש את התנועה שלו. כך, כשהלוח קורא getMoves, הוא לא "
                       "צריך לדעת איזה כלי זה — זה פולימורפיזם, וזה חסך לי המון קוד כפול.", "he"),
                      ("\"", "he")], size=12, space_after=8)

    heading(doc, "5.5 שאלות ותשובות — OOP", level=2)
    qa_list(doc, [
        ("מה זה מחלקה מופשטת (abstract)?",
         "מחלקה שאי אפשר ליצור ממנה אובייקט ישירות, ויכולה להכריח את היורשים לממש "
         "מתודות מסוימות. אצלי Piece היא abstract עם getMoves מופשטת."),
        ("מה זה ירושה ואיך ניצלת אותה?",
         "מנגנון שבו מחלקה מקבלת תכונות ופעולות מאב. כל הכלים יורשים מ-Piece ומקבלים "
         "ממנה את color, hasMoved, slideMoves ו-jumpMoves."),
        ("מה זה פולימורפיזם? תן דוגמה.",
         "אותה קריאה (getMoves) מתנהגת אחרת בכל כלי. הלוח קורא לה בלי לדעת את סוג הכלי."),
        ("מה זה כימוס?",
         "הסתרת פרטים פנימיים מאחורי ממשק נקי. ב-Board הלוח grid פרטי, וחושפים "
         "רק פעולות כמו applyMove ו-legalMoves."),
    ])


def write_protocol(doc):
    heading(doc, "6.1 הרעיון: שפה משותפת", level=2)
    he_paragraph(doc,
                 "כדי שהשרת והלקוח יבינו זה את זה, הם מדברים ב'שפה' מוסכמת — "
                 "פרוטוקול. כל הודעה היא טקסט בפורמט JSON, שורה אחת לכל הודעה. "
                 "לכל הודעה יש שדה type שאומר 'מה זה' (למשל login, move), ושדות "
                 "נוספים לפי הצורך.")
    he_paragraph(doc, "דוגמאות להודעות לקוח → שרת:")
    add_table(doc, ["type", "תיאור", "שדות"], [
        ("register / login", "הרשמה / התחברות", "username, password"),
        ("play_human", "הצטרפות לתור מול אדם", "(אין)"),
        ("play_ai", "משחק נגד מחשב", "level"),
        ("get_moves", "בקשת מהלכים חוקיים לכלי", "col, row"),
        ("move", "ביצוע מהלך", "from_col, from_row, to_col, to_row"),
        ("promote", "בחירת כלי לקידום חייל", "choice"),
        ("resign", "כניעה", "(אין)"),
    ])
    he_paragraph(doc, "דוגמאות להודעות שרת → לקוח:")
    add_table(doc, ["type", "תיאור"], [
        ("auth_result", "תוצאת הרשמה/התחברות (ok, message)"),
        ("info / error", "הודעת מידע / שגיאה"),
        ("game_start", "המשחק התחיל (color, opponent)"),
        ("moves", "רשימת מהלכים חוקיים לכלי"),
        ("state", "מצב הלוח המעודכן (board, turn, status, winner)"),
        ("choose_promotion", "בקשה לבחור כלי לקידום"),
    ])

    heading(doc, "6.2 איך זה מומש — JSONUtil", level=2)
    he_paragraph(doc,
                 "לא השתמשתי בשום ספרייה חיצונית — כתבתי מקודד ומפענח JSON משלי. "
                 "encode מקבל מפה (Map) של שדות והופך אותה לטקסט; decode עושה את "
                 "ההפך — מטקסט למפה:")
    code_block(doc, [
        "// JSONUtil.java — בניית טקסט JSON ממפה",
        "public static String encode(Map<String, Object> map) {",
        "    StringBuilder sb = new StringBuilder(\"{\");",
        "    boolean first = true;",
        "    for (Map.Entry<String, Object> e : map.entrySet()) {",
        "        if (!first) sb.append(\",\");",
        "        sb.append(\"\\\"\").append(e.getKey()).append(\"\\\":\");",
        "        sb.append(encodeValue(e.getValue()));",
        "        first = false;",
        "    }",
        "    return sb.append(\"}\").toString();",
        "}",
    ])
    he_paragraph(doc,
                 "בצד המקבל, ClientHandler קורא שורה, מפענח אותה ל-Map, ולפי השדה "
                 "type מפנה לטיפול הנכון (switch):")
    code_block(doc, [
        "// ClientHandler.java",
        "while ((line = in.readLine()) != null) {",
        "    handleMessage(JSONUtil.decode(line));",
        "}",
        "...",
        "switch (type) {",
        "    case \"login\":  handleAuth(msg, true); break;",
        "    case \"move\":   handleMove(msg);       break;",
        "    case \"play_ai\": handlePlayAI(msg);     break;",
        "    // ...",
        "}",
    ])

    heading(doc, "6.3 דוגמת זרימה מלאה", level=2)
    code_block(doc, [
        'C -> S  {"type":"login","username":"alice","password":"1234"}',
        'S -> C  {"type":"auth_result","ok":true,"username":"alice"}',
        "",
        'C -> S  {"type":"play_ai","level":2}',
        'S -> C  {"type":"game_start","color":"white","opponent":"AI"}',
        'S -> C  {"type":"state","board":[...],"turn":"white",...}',
        "",
        'C -> S  {"type":"move","from_col":4,"from_row":1,"to_col":4,"to_row":3}',
        'S -> C  {"type":"state","board":[...],"turn":"black","last_move":[4,1,4,3]}',
        'S -> C  {"type":"state","board":[...],"turn":"white",...}   // מהלך המחשב',
    ])

    heading(doc, "6.4 שאלות ותשובות — פרוטוקול", level=2)
    qa_list(doc, [
        ("למה בחרת JSON?",
         "כי הוא קריא לבני אדם, קל לפענוח, וקל לבדיקה ולניפוי שגיאות. חלופה בינארית "
         "קומפקטית יותר אך קשה לקריאה."),
        ("איך השרת יודע מה לעשות עם הודעה?",
         "לפי השדה type. יש switch ב-handleMessage שמפנה כל סוג הודעה לפונקציה "
         "המתאימה."),
        ("איפה מפענחים את הטקסט?",
         "ב-JSONUtil.decode — שממיר את שורת הטקסט ל-Map של שדות וערכים."),
    ])


def write_server_structure(doc):
    he_paragraph(doc,
                 "חשוב להגנה להראות שאתה יודע מי אחראי על מה. חלוקת התפקידים בשרת:")
    add_table(doc, ["מחלקה", "אחריות"], [
        ("ChessServer", "מאזין לחיבורים (accept) ופותח נימה לכל לקוח"),
        ("ClientHandler", "מייצג לקוח אחד — קורא הודעות, מפענח, ומגיב. מחזיק את ה-Socket"),
        ("Matchmaker", "תור FIFO שמזווג שני שחקנים שמחכים למשחק מול אדם"),
        ("Game", "משחק בודד — מחזיק Board ושני ClientHandler-ים, מבצע מהלכים, ומשדר מצב"),
        ("Board", "מודל הלוח: כלים, חוקי תנועה, שח/מט/פט, קידום, הצרחה, en passant"),
        ("UserAuth", "הרשמה/התחברות וערבול סיסמאות + שמירה ב-users.json"),
    ])
    he_paragraph(doc,
                 "שים לב לזרימה: ClientHandler מקבל הודעת move, קורא ל-Game.applyMove "
                 "(שהיא synchronized), ו-Game מבקש מ-Board לבדוק חוקיות ולבצע. בסוף "
                 "Game משדר את המצב החדש לשני הלקוחות עם broadcastState():")
    code_block(doc, [
        "// Game.java — שידור המצב לשני הצדדים",
        "private void broadcastState() {",
        "    String json = buildState();",
        "    white.send(json);",
        "    if (black != null) black.send(json);   // ב-AI אין 'שחור' אנושי",
        "}",
    ])
    he_paragraph(doc,
                 "אחרי כל מהלך אנושי, אם זה משחק נגד המחשב ועכשיו תור השחור — "
                 "Game מבקש מ-ChessAI מהלך ומבצע אותו אוטומטית:")
    code_block(doc, [
        "// Game.java — afterMove()",
        "if (ai != null && board.turn == Color.BLACK) {",
        "    int[] move = ai.pickMove(board);",
        "    if (move != null) {",
        "        board.applyMove(move[0], move[1], move[2], move[3]);",
        "        broadcastState();",
        "    }",
        "}",
    ])


def write_ai(doc):
    heading(doc, "8.1 הרעיון: חמדנות (greedy)", level=2)
    he_paragraph(doc,
                 "הבינה המלאכותית פשוטה בכוונה — מתאימה להיקף של פרויקט 5 יחידות. "
                 "ברמה 1 היא בוחרת מהלך אקראי מבין החוקיים. ברמה גבוהה יותר היא "
                 "'חמדנית': לכל מהלך חוקי היא מחשבת ניקוד, ובוחרת את בעל הניקוד הגבוה ביותר.")
    he_paragraph(doc, "הניקוד מורכב משני חלקים:")
    he_bullet(doc, "ערך הכלי שנאכל (אם נאכל) × 10 — עדיף לאכול מלכה מאשר חייל.")
    he_bullet(doc, "בונוס קטן על התקדמות למרכז הלוח.")
    code_block(doc, [
        "// ChessAI.java — חישוב ניקוד למהלך",
        "private int scoreMove(Board board, int[] m) {",
        "    int score = 0;",
        "    Piece target = board.getPiece(m[2], m[3]);",
        "    if (target != null) score += target.getValue() * 10;   // אכילה",
        "    int dc = m[2] - 3, dr = m[3] - 3;",
        "    int centerDist = (int) Math.sqrt(dc*dc + dr*dr);",
        "    score += Math.max(0, 4 - centerDist);                  // קרבה למרכז",
        "    return score;",
        "}",
    ])
    he_paragraph(doc,
                 "המחשב אוסף את כל המהלכים החוקיים שלו (על כל הכלים), מנקד כל אחד, "
                 "ובוחר את הטוב ביותר (עם שובר-שוויון אקראי):")
    code_block(doc, [
        "// ChessAI.java — בחירת המהלך הטוב ביותר",
        "for (int[] m : moves) {",
        "    int score = scoreMove(board, m);",
        "    if (score > bestScore) { bestScore = score; bestMoves.clear(); bestMoves.add(m); }",
        "    else if (score == bestScore) { bestMoves.add(m); }",
        "}",
        "return bestMoves.get(rng.nextInt(bestMoves.size()));",
    ])
    callout(doc, [
        ("חלופה שנדחתה: ", "bold"),
        ("Minimax עם Alpha-Beta מסתכל כמה מהלכים קדימה והוא חזק יותר, אבל מורכב "
         "מדי לרמת הפרויקט. בחרתי בחמדנות כי היא פשוטה להבנה ומנצחת מהלכים אקראיים.", "he"),
    ])

    heading(doc, "8.2 שאלות ותשובות — AI", level=2)
    qa_list(doc, [
        ("איך המחשב בוחר מהלך?",
         "אוסף את כל המהלכים החוקיים, מנקד כל אחד (אכילה + קרבה למרכז), ובוחר את "
         "בעל הניקוד הגבוה ביותר."),
        ("למה לא Minimax?",
         "כי הוא מורכב מדי להיקף של 5 יחידות. חמדנות פשוטה מספיקה כדי להדגים בינה "
         "מלאכותית ולנצח משחק אקראי."),
    ])


def write_chess_rules(doc):
    heading(doc, "9.1 בדיקת חוקיות מהלך", level=2)
    he_paragraph(doc,
                 "מהלך חוקי הוא לא רק מהלך אפשרי לפי הכלי, אלא גם כזה שלא משאיר את "
                 "המלך שלי בשח. כדי לבדוק, מבצעים את המהלך באופן זמני, בודקים אם "
                 "המלך מאוים, ומחזירים את הלוח לאחור:")
    code_block(doc, [
        "// Board.java — moveLeavesKingInCheck()",
        "Piece saved = grid[tc][tr];",
        "grid[tc][tr] = grid[fc][fr];   // בצע זמנית",
        "grid[fc][fr] = null;",
        "boolean inCheck = isInCheck(turn);   // בדוק שח",
        "grid[fc][fr] = grid[tc][tr];   // החזר לאחור",
        "grid[tc][tr] = saved;",
        "return inCheck;",
    ])
    he_paragraph(doc,
                 "legalMoves מסנן את כל המהלכים שמשאירים את המלך בשח, ומוסיף הצרחה "
                 "אם מדובר במלך. כך הלקוח מקבל רק מהלכים שבאמת חוקיים.")

    heading(doc, "9.2 חוקים מיוחדים שמומשו", level=2)
    he_bullet(doc, "קידום חייל — חייל שמגיע לשורה האחרונה: השחקן בוחר מלכה/צריח/רץ/פרש (הודעת choose_promotion).")
    he_bullet(doc, "הכאה דרך הילוכו (en passant) — נשמר בשדה enPassant אחרי זינוק חייל שתי משבצות.")
    he_bullet(doc, "הצרחה (castling) — נבדקים חמשת תנאי ה-FIDE: המלך והצריח לא זזו, הנתיב ריק, המלך לא בשח ולא עובר דרך משבצת מאוימת.")
    he_bullet(doc, "שח / מט / פט — מחושבים אחרי כל מהלך ומעדכנים את status.")

    heading(doc, "9.3 שאלות ותשובות — חוקי שחמט", level=2)
    qa_list(doc, [
        ("איך אתה בודק שמהלך לא משאיר את המלך בשח?",
         "מבצע את המהלך זמנית על הלוח, בודק isInCheck, ואז מחזיר את הלוח למצב הקודם."),
        ("איך מומשה ההצרחה?",
         "במתודה addCastlingMoves שבודקת את כל תנאי ה-FIDE לפני שהיא מוסיפה את "
         "מהלך ההצרחה לרשימת המהלכים החוקיים."),
        ("מה ההבדל בין מט לפט?",
         "בשניהם אין מהלך חוקי. במט המלך מאוים (הפסד), בפט המלך לא מאוים (תיקו)."),
    ])


def write_gui(doc):
    heading(doc, "10.1 מה זה Swing ולמה השתמשתי בו", level=2)
    he_paragraph(doc,
                 "Swing היא ספריית הממשק הגרפי המובנית של Java (אין צורך בהתקנה). "
                 "איתה בונים חלונות, כפתורים, שדות טקסט ואזורי ציור. בחרתי בה כי היא "
                 "חלק מ-Java SE, פשוטה, ורצה על כל מערכת הפעלה.")
    he_paragraph(doc, "מושגי היסוד של Swing בפרויקט:")
    he_bullet(doc, "JFrame — החלון הראשי. המחלקה ChessGUI יורשת מ-JFrame.")
    he_bullet(doc, "JPanel — לוח שמכיל רכיבים (כפתורים, תוויות) או משמש לציור.")
    he_bullet(doc, "JButton / JTextField / JPasswordField / JLabel — כפתורים, שדות וטקסט.")
    he_bullet(doc, "CardLayout — מאפשר להחליף בין כמה 'מסכים' באותו חלון, כמו ערימת קלפים.")

    heading(doc, "10.2 שלושה מסכים עם CardLayout", level=2)
    he_paragraph(doc,
                 "ב-ChessGUI יש שלושה מסכים: התחברות (login), לובי (lobby) ולוח "
                 "המשחק (game). כולם נטענים מראש ל-CardLayout, ובכל רגע מציגים אחד "
                 "מהם עם cards.show(...):")
    code_block(doc, [
        "// ChessGUI.java — שלושת המסכים",
        "private final CardLayout cards = new CardLayout();",
        "private final JPanel root = new JPanel(cards);",
        "...",
        "root.add(buildLoginPanel(), \"login\");",
        "root.add(buildLobbyPanel(), \"lobby\");",
        "root.add(buildGamePanel(),  \"game\");",
        "cards.show(root, \"login\");   // מתחילים במסך ההתחברות",
    ])
    he_paragraph(doc,
                 "מעבר בין מסכים קורה בתגובה לאירועים — למשל אחרי התחברות מוצלחת "
                 "עוברים ללובי: cards.show(root, \"lobby\").")

    heading(doc, "10.3 כפתורים ששולחים הודעות (Action Listeners)", level=2)
    he_paragraph(doc,
                 "כל כפתור מקבל 'מאזין' (ActionListener) שרץ בלחיצה. הלחיצה בונה "
                 "הודעת JSON ושולחת אותה לשרת דרך NetworkClient. כך לחיצה על 'Play "
                 "vs AI' שולחת play_ai:")
    code_block(doc, [
        "// ChessGUI.java",
        "btnAI2.addActionListener(e -> playAI(2));",
        "...",
        "private void playAI(int level) {",
        "    Map<String, Object> msg = new LinkedHashMap<>();",
        "    msg.put(\"type\", \"play_ai\");",
        "    msg.put(\"level\", level);",
        "    net.send(msg);            // שליחה לשרת",
        "}",
    ])

    heading(doc, "10.4 איך הודעות מהשרת מעדכנות את המסך (Timer)", level=2)
    he_paragraph(doc,
                 "נזכיר: נימת הרקע של NetworkClient שמה הודעות נכנסות בתור (inbox). "
                 "אבל אסור לעדכן רכיבי Swing מנימה אחרת. לכן הוספתי Timer של Swing "
                 "שמתעורר כל 100 מילי-שניות, שולף את כל ההודעות מהתור, ומעדכן את "
                 "המסך — הכול בנימת ה-UI הראשית:")
    code_block(doc, [
        "// ChessGUI.java",
        "pollTimer = new Timer(100, e -> pollMessages());",
        "pollTimer.start();",
        "...",
        "private void pollMessages() {",
        "    Map<String, Object> msg;",
        "    while ((msg = net.poll()) != null) {",
        "        handleMessage(msg);   // לפי type — מעדכן מסך מתאים",
        "    }",
        "}",
    ])
    callout(doc, [
        ("נקודה חשובה להגנה: ", "bold"),
        ("הפרדתי בין נימת הרשת (שמקבלת הודעות) לבין נימת ה-UI (שמציירת). ה-Timer "
         "הוא 'הגשר' הבטוח ביניהן — הוא רץ בנימת ה-UI, ולכן מותר לו לעדכן את המסך.", "he"),
    ])
    he_paragraph(doc, "handleMessage בצד הלקוח מפנה כל סוג הודעה לטיפול המתאים:")
    code_block(doc, [
        "// ChessGUI.java",
        "switch (type) {",
        "    case \"auth_result\":      handleAuthResult(msg); break;",
        "    case \"game_start\":       handleGameStart(msg);  break;",
        "    case \"state\":            handleState(msg);      break;",
        "    case \"moves\":            handleMoves(msg);      break;",
        "    case \"choose_promotion\": handleChoosePromotion(); break;",
        "}",
    ])

    heading(doc, "10.5 ציור הלוח — paintComponent", level=2)
    he_paragraph(doc,
                 "הלוח עצמו הוא מחלקה פנימית BoardPanel שיורשת מ-JPanel. כדי לצייר "
                 "מותאם אישית דורסים את paintComponent — Swing קורא לה בכל פעם שצריך "
                 "לצייר מחדש (אחרי repaint()). הציור נעשה בשלושה שלבים:")
    code_block(doc, [
        "// ChessGUI.BoardPanel",
        "protected void paintComponent(Graphics g) {",
        "    super.paintComponent(g);",
        "    Graphics2D g2 = (Graphics2D) g;",
        "    drawBoard(g2);                       // 1. משבצות + הדגשות",
        "    if (boardData != null) drawPieces(g2);// 2. הכלים",
        "    drawHints(g2);                       // 3. נקודות המהלכים הירוקות",
        "}",
    ])
    he_bullet(doc, "drawBoard — מצייר 64 משבצות בשני צבעים, מדגיש את הכלי הנבחר ואת המהלך האחרון, ומוסיף אותיות/מספרים בקצוות (a-h, 1-8).")
    he_bullet(doc, "drawPieces — מצייר כל כלי בעזרת תו יוניקוד (♔ ♕ ♖ ♗ ♘ ♙ ...) לפי הקוד שמגיע מהשרת.")
    he_bullet(doc, "drawHints — מצייר את 'הנקודות הירוקות' של המהלכים החוקיים.")

    heading(doc, "10.6 הדגשת המהלכים החוקיים (ה'נקודות הירוקות')", level=2)
    he_paragraph(doc,
                 "כשלוחצים על כלי, הלקוח שולח get_moves, השרת מחזיר moves, והלקוח "
                 "שומר אותם ומצייר: עיגול ירוק מלא למשבצת ריקה, וטבעת ירוקה סביב כלי "
                 "יריב שאפשר לאכול:")
    code_block(doc, [
        "// ChessGUI.BoardPanel — drawHints()",
        "for (int[] m : legalMoveHints) {",
        "    String cell = boardData[m[1]][m[0]];",
        "    boolean isCapture = cell != null && !cell.equals(\".\");",
        "    if (isCapture) {",
        "        g.drawOval(...);   // טבעת — אפשר לאכול",
        "    } else {",
        "        g.fillOval(...);   // עיגול מלא — משבצת פנויה",
        "    }",
        "}",
    ])

    heading(doc, "10.7 לחיצה על הלוח — מהקליק למהלך", level=2)
    he_paragraph(doc,
                 "MouseListener תופס לחיצות. הלחיצה הראשונה בוחרת כלי (ושולחת "
                 "get_moves), הלחיצה השנייה שולחת את המהלך עצמו. כי יש להמיר פיקסלים "
                 "למשבצת לוח, ובגלל שהלוח מתהפך כשמשחקים בשחור — יש מתודות המרה:")
    code_block(doc, [
        "// ChessGUI.java — onBoardClick (מקוצר)",
        "if (selectedCell == null) {            // בחירת כלי",
        "    selectedCell = new int[]{col, rank};",
        "    req.put(\"type\", \"get_moves\");      // לקבל מהלכים חוקיים",
        "    net.send(req);",
        "} else {                                // ביצוע מהלך",
        "    msg.put(\"type\", \"move\");",
        "    msg.put(\"from_col\", selectedCell[0]);",
        "    msg.put(\"from_row\", selectedCell[1]);",
        "    msg.put(\"to_col\", col);",
        "    msg.put(\"to_row\", rank);",
        "    net.send(msg);",
        "}",
    ])
    he_paragraph(doc,
                 "היפוך הלוח לשחקן השחור נעשה ב-colToScreenX/rankToScreenY (לציור) "
                 "ו-screenToCol/screenToRank (ללחיצה), כך ששני השחקנים תמיד רואים "
                 "את הכלים שלהם מלמטה.")

    heading(doc, "10.8 דיאלוג קידום החייל", level=2)
    he_paragraph(doc,
                 "כשמגיעה הודעת choose_promotion מהשרת, הלקוח פותח חלון בחירה "
                 "(JOptionPane) ושולח בחזרה את הבחירה (Q/R/B/N):")
    code_block(doc, [
        "// ChessGUI.java",
        "String[] options = {\"Queen\", \"Rook\", \"Bishop\", \"Knight\"};",
        "int choice = JOptionPane.showOptionDialog(... options ...);",
        "// ממירים לאות ושולחים promote לשרת",
    ])

    heading(doc, "10.9 איך להגיד את זה בהגנה", level=2)
    he_segments(doc, [("\"", "he"),
                      ("בניתי את הממשק ב-Swing עם שלושה מסכים ב-CardLayout. כל כפתור "
                       "שולח הודעת JSON לשרת. הודעות מהשרת נכנסות לתור בנימת רקע, "
                       "ו-Timer של Swing שולף אותן כל 100 מ\"ש ומעדכן את המסך בבטחה. "
                       "את הלוח ציירתי לבד ב-paintComponent, כולל הדגשת המהלכים "
                       "החוקיים בנקודות ירוקות.", "he"),
                      ("\"", "he")], size=12, space_after=8)

    heading(doc, "10.10 שאלות ותשובות — GUI", level=2)
    qa_list(doc, [
        ("מה זה Swing?",
         "ספריית הממשק הגרפי המובנית של Java לבניית חלונות, כפתורים ואזורי ציור."),
        ("איך עברת בין מסכים?",
         "עם CardLayout — שלושה JPanel (login, lobby, game) שמתחלפים עם cards.show()."),
        ("איך המסך מתעדכן כשמגיעה הודעה מהשרת?",
         "נימת רקע שמה הודעות בתור, ו-Timer של Swing שולף אותן כל 100 מ\"ש ומעדכן "
         "את הרכיבים בנימת ה-UI (אסור לעדכן Swing מנימה אחרת)."),
        ("איך ציירת את הלוח?",
         "דרסתי את paintComponent ב-JPanel פנימי, וציירתי משבצות, כלים (תווי יוניקוד) "
         "והדגשות בעזרת Graphics2D."),
        ("איך עובדת בחירת מהלך בעכבר?",
         "לחיצה ראשונה בוחרת כלי ושולחת get_moves; לחיצה שנייה שולחת move. ממירים "
         "פיקסלים למשבצת, ומתחשבים בהיפוך הלוח לשחקן השחור."),
    ])


def write_glossary(doc):
    add_table(doc, ["מונח", "פירוש קצר"], [
        ("נימה (Thread)", "נתיב ריצה בתוך תוכנית; כל הנימות חולקות זיכרון"),
        ("synchronized", "מילת מפתח שמרשה רק לנימה אחת להיכנס למתודה בכל רגע"),
        ("Socket", "קצה החיבור שדרכו שולחים ומקבלים נתונים"),
        ("ServerSocket", "אובייקט בצד השרת שמאזין לחיבורים נכנסים"),
        ("TCP", "פרוטוקול תעבורה אמין ומסודר"),
        ("פורט (Port)", "מספר שמזהה תוכנית מסוימת במחשב (אצלי 5555)"),
        ("ערבול (Hash)", "פונקציה חד-כיוונית שממירה קלט לפלט בגודל קבוע"),
        ("SHA-256", "אלגוריתם ערבול סטנדרטי"),
        ("salt", "מחרוזת אקראית שמתווספת לסיסמה לפני הערבול"),
        ("Runnable", "ממשק עם מתודת run() — הקוד שנימה מריצה"),
        ("פולימורפיזם", "אותה קריאה מתנהגת אחרת לפי סוג האובייקט"),
        ("abstract", "מחלקה/מתודה שאי אפשר להשתמש בה ישירות; חייבים לממש"),
        ("JSON", "פורמט טקסט לייצוג נתונים, משמש להודעות בין שרת ללקוח"),
        ("Swing", "ספריית הממשק הגרפי המובנית של Java"),
        ("JFrame / JPanel", "חלון / לוח רכיבים או אזור ציור ב-Swing"),
        ("CardLayout", "פריסה שמחליפה בין כמה מסכים באותו חלון"),
        ("paintComponent", "המתודה שדורסים כדי לצייר ציור מותאם ב-JPanel"),
    ])


def write_qa_bank(doc):
    he_paragraph(doc,
                 "שאלות שכדאי לתרגל לפני ההגנה — כסה את התשובה ונסה לענות בקול.")
    qa_list(doc, [
        ("ספר במשפט מה הפרויקט עושה.",
         "משחק שחמט בארכיטקטורת שרת/לקוח: שני שחקנים מתחברים דרך הרשת, נרשמים "
         "עם סיסמה מאובטחת, ומשחקים זה נגד זה או נגד המחשב."),
        ("איפה בא לידי ביטוי ריבוי נימות?",
         "השרת פותח נימה לכל לקוח (new Thread(handler).start()), כך שכמה משחקים "
         "רצים במקביל."),
        ("איך הגנת על מידע משותף בין נימות?",
         "מתודות המשחק, התור והאימות מסומנות synchronized."),
        ("איך אבטחת את הסיסמאות?",
         "לא שמרתי סיסמה גלויה — רק salt אקראי וערבול SHA-256 של (salt+סיסמה)."),
        ("מה ההבדל בין ערבול להצפנה?",
         "ערבול חד-כיווני (לא ניתן לפענוח), הצפנה דו-כיוונית. לסיסמאות משתמשים בערבול."),
        ("למה TCP ולא UDP?",
         "TCP אמין ומסודר — חשוב כדי שמהלך לא יאבד או יגיע בסדר הלא נכון."),
        ("איפה השתמשת בירושה ופולימורפיזם?",
         "כל הכלים יורשים ממחלקת האב המופשטת Piece ומממשים getMoves שלהם; הלוח "
         "קורא getMoves בלי לדעת את סוג הכלי."),
        ("מי בודק שמהלך חוקי — הלקוח או השרת?",
         "השרת. הוא הסמכות — בודק כל מהלך לפני ביצוע, כדי שאי אפשר יהיה לרמות."),
        ("איך עובד פרוטוקול ההודעות?",
         "כל הודעה היא JSON עם שדה type. הצד המקבל מפענח עם JSONUtil ומפנה לפי "
         "ה-type ב-switch."),
        ("מה היית משפר אם היה לך עוד זמן?",
         "AI חזק יותר (Minimax), שעון לכל שחקן, שמירת היסטוריית משחקים, וכלל "
         "חמישים המהלכים."),
    ])


def qa_list(doc, pairs):
    for q, a in pairs:
        he_segments(doc, [("ש: ", "bold"), (q, "he")], size=12)
        he_segments(doc, [("ת: ", "accent"), (a, "he")], size=12, space_after=8)


if __name__ == "__main__":
    build_document()
