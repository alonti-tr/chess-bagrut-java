from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_helpers import (
    ACCENT, ACCENT_DARK, DARK, MUTED,
    add_page_numbers, add_table, callout, code_block, he_bullet,
    he_paragraph, he_segments, heading, new_document, page_break,
    set_update_fields_on_open, tip_box, word_postprocess_rtl, _set_bidi,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = REPO_ROOT / "תיק-פרויקט.docx"

SOURCE_FILES = [
    "src/chess/Color.java",
    "src/chess/pieces/Piece.java",
    "src/chess/pieces/Pawn.java",
    "src/chess/pieces/Knight.java",
    "src/chess/pieces/Bishop.java",
    "src/chess/pieces/Rook.java",
    "src/chess/pieces/Queen.java",
    "src/chess/pieces/King.java",
    "src/chess/Board.java",
    "src/chess/ChessAI.java",
    "src/chess/UserAuth.java",
    "src/chess/JSONUtil.java",
    "src/chess/server/Matchmaker.java",
    "src/chess/server/Game.java",
    "src/chess/server/ClientHandler.java",
    "src/chess/server/ChessServer.java",
    "src/chess/client/NetworkClient.java",
    "src/chess/client/ChessGUI.java",
    "src/chess/Main.java",
]


def build_document() -> None:
    doc = new_document()
    add_page_numbers(doc)

    write_cover(doc)
    page_break(doc)

    write_toc(doc)
    page_break(doc)

    heading(doc, "1. מבוא — ייזום ואפיון", level=1)
    write_intro(doc)
    page_break(doc)

    heading(doc, "2. תיאור תחום הידע — ניתוח", level=1)
    write_knowledge_domain(doc)
    page_break(doc)

    heading(doc, "3. ארכיטקטורה (העיצוב)", level=1)
    write_architecture(doc)
    page_break(doc)

    heading(doc, "4. מימוש הפרויקט — הקוד", level=1)
    write_implementation(doc)
    page_break(doc)

    heading(doc, "5. מדריך למשתמש", level=1)
    write_user_guide(doc)
    page_break(doc)

    heading(doc, "6. סיכום אישי ורפלקציה", level=1)
    write_reflection(doc)
    page_break(doc)

    heading(doc, "7. ביבליוגרפיה — סקר ספרות", level=1)
    write_bibliography(doc)
    page_break(doc)

    heading(doc, "נספח א' — תדפיס הקוד המקור המלא", level=1)
    write_source_appendix(doc)
    page_break(doc)

    heading(doc, "נספח ב' — טבלת הודעות הפרוטוקול", level=1)
    write_protocol_appendix(doc)
    page_break(doc)

    heading(doc, "נספח ג' — שאלות תיאורטיות אפשריות לבחינה", level=1)
    write_qa_appendix(doc)

    set_update_fields_on_open(doc)
    doc.save(str(OUT_DOCX))
    word_postprocess_rtl(OUT_DOCX)
    print(f"wrote portfolio docx ({OUT_DOCX.stat().st_size:,} bytes)")


def write_cover(doc):
    he_paragraph(doc, "", space_before=60)
    he_paragraph(doc, "תיק פרויקט", size=26, bold=True, color=ACCENT_DARK,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    he_paragraph(doc, "משחק שחמט מאובטח בין שני משתתפים (Java)", size=28, bold=True, color=ACCENT,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    he_paragraph(doc, "מערכת שרת/לקוח מרובת משתמשים, עם תור המתנה, משחק נגד מחשב, והצפנת סיסמאות",
                 size=13, color=MUTED, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    he_paragraph(doc, "פרויקט גמר במדעי המחשב", size=12, color=DARK,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    he_paragraph(doc, "5 יחידות לימוד — מימוש בשפת Java", size=12, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    fields = [
        ("שם בית הספר", "הכפר הירוק"),
        ("סמל מוסד", "_____________________"),
        ("שם התלמיד", "לירן תדהר"),
        ('מספר ת"ז', "218046456"),
        ("שם המורה המנחה", "יהודה אור"),
        ("שם נושא הפרויקט", "משחק שחמט מאובטח בין שני משתתפים"),
        ("תאריך הגשה", "7.6.2026"),
    ]
    add_table(doc, ["שדה", "ערך"], fields, col_widths=[5.5, 11.0])


def write_toc(doc):
    heading(doc, "תוכן עניינים", level=1)
    he_paragraph(doc,
                 "להלן תוכן העניינים האוטומטי של Word. כדי לעדכן אותו לאחר שינויים: "
                 "לחץ קליק־ימני על הטבלה ובחר 'Update Field' (או F9).",
                 italic=True, color=MUTED)

    p = doc.add_paragraph()
    _set_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\p " "'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = p.add_run()._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

    he_paragraph(doc, "תוכן עניינים — לחץ F9 ב-Word כדי לעדכן", italic=True, color=MUTED)


def write_intro(doc):
    heading(doc, "1.1 תיאור הרעיון והמוטיבציה", level=2)
    he_paragraph(doc,
                 "הפרויקט מממש משחק שחמט לשני משתתפים בארכיטקטורת שרת/לקוח, בשפת Java. "
                 "שני שחקנים מתחברים לשרת מרכזי, נרשמים או מתחברים עם שם משתמש וסיסמה, "
                 "ממתינים בתור עד שמצטרף יריב, ואז משחקים זה נגד זה בזמן אמת דרך "
                 "תקשורת רשת מבוססת TCP. בנוסף, שחקן יחיד יכול לשחק נגד מחשב (בינה "
                 "מלאכותית פשוטה). הסיסמאות נשמרות בקובץ בצורה מוצפנת באמצעות "
                 "פונקציית גיבוב (hash) מסוג SHA-256 עם salt אקראי לכל משתמש.")
    he_paragraph(doc,
                 "המוטיבציה: שחמט הוא משחק עם חוקים מוגדרים היטב, ולכן הוא דוגמה "
                 "טובה להדגמת מגוון נושאים בפרויקט אחד — תכנות מונחה עצמים (מחלקות "
                 "וירושה), תקשורת מבוססת סוקטים, ריבוי תהליכונים, פרוטוקול הודעות "
                 "פשוט, שמירת נתונים בקובץ, גיבוב סיסמאות, וממשק משתמש גרפי ב-Swing.")

    heading(doc, "1.2 ייזום: זיהוי הצורך", level=2)
    he_paragraph(doc,
                 "רציתי לבנות משחק שאפשר לשחק בו מול חבר על שני מחשבים שונים, ולא "
                 "רק מול עצמך על אותו מסך. זה דרש ללמוד איך מחשבים מדברים ביניהם "
                 "ברשת. בנוסף רציתי שהמערכת תהיה בטוחה במידה סבירה — שהסיסמאות לא "
                 "יישמרו בטקסט גלוי בקובץ, כי זו טעות אבטחה נפוצה.")

    heading(doc, "1.3 אפיון פונקציונלי — דרישות המשתמש", level=2)
    add_table(doc, ["#", "דרישה", "תיאור"], [
        ("FR-1", "הרשמה", "משתמש יכול להירשם עם שם משתמש וסיסמה"),
        ("FR-2", "התחברות", "משתמש קיים יכול להתחבר; מתקבלת שגיאה אם הפרטים שגויים"),
        ("FR-3", "תור משחק", "שחקן מחובר יכול להצטרף לתור ולהמתין עד שמצטרף יריב"),
        ("FR-4", "משחק אדם מול אדם", "שני שחקנים משחקים זה נגד זה; הראשון שהמתין מקבל לבן"),
        ("FR-5", "משחק נגד מחשב", "שחקן יכול לבחור לשחק נגד המחשב"),
        ("FR-6", "מהלכים חוקיים בלבד", "השרת בודק שכל מהלך חוקי לפי חוקי השחמט שמומשו"),
        ("FR-7", "קידום חייל בבחירה", "חייל שמגיע לשורה האחרונה — השחקן בוחר מלכה/צריח/רץ/פרש"),
        ("FR-8", "הכאה דרך הילוכו", "נתמך en passant לפי החוק הרשמי"),
        ("FR-9", "סיום משחק", "מט / פט / כניעה / ניתוק → המשחק נגמר והודעה נשלחת לשני הצדדים"),
        ("FR-10", "כניעה", "שחקן יכול להיכנע באמצע משחק; היריב מנצח"),
    ])

    heading(doc, "1.4 אפיון לא-פונקציונלי", level=2)
    add_table(doc, ["#", "תכונה", "יעד"], [
        ("NFR-1", "אבטחת סיסמאות", "סיסמאות נשמרות מגובבות בלבד (SHA-256 + salt). אין סיסמה גלויה בקובץ users.json"),
        ("NFR-2", "יציבות שרת", "ניתוק של לקוח אחד אינו מפיל את השרת או משחקים אחרים"),
        ("NFR-3", "ריבוי משחקים", "השרת תומך בכמה משחקים במקביל בעזרת תהליכון לכל לקוח"),
        ("NFR-4", "פשטות", "הקוד כתוב ב-Java SE בלבד, בלי ספריות חיצוניות, וקריא ללומד"),
        ("NFR-5", "ניידות", "רץ על Windows, Linux ו-macOS עם JDK 17 ומעלה"),
    ])

    heading(doc, "1.5 לוח זמנים — אבני דרך", level=2)
    add_table(doc, ["שלב", "תוצר", "מצב"], [
        ("שלב 1 — מחלקות הכלים והלוח", "חבילת pieces + Board.java (ירושה, חוקי תנועה, שח/מט)", "הושלם"),
        ("שלב 2 — אימות משתמשים", "UserAuth.java (SHA-256 + salt + users.json)", "הושלם"),
        ("שלב 3 — שרת ולקוח", "חבילת server + client (TCP, תהליכונים, תור)", "הושלם"),
        ("שלב 4 — בינה מלאכותית", "ChessAI.java (חמדנית) — בונוס", "הושלם"),
        ("שלב 5 — ממשק גרפי", "ChessGUI.java עם Swing (3 מסכים)", "הושלם"),
        ("שלב 6 — תיק פרויקט ומחוון", "תיק + מחוון", "הושלם"),
    ])

    heading(doc, "1.6 ניהול סיכונים", level=2)
    add_table(doc, ["סיכון", "הסתברות", "חומרה", "מיטיגציה"], [
        ("נפילת שרת בגלל הודעה פגומה", "בינונית", "גבוהה",
         "כל טיפול בהודעה עטוף ב-try/catch; שגיאה הופכת להודעת שגיאה ולא מפילה את ה-Thread"),
        ("שני תהליכונים ניגשים לאותו משחק", "בינונית", "בינונית",
         "מתודות המשחק מסומנות synchronized — רק תהליכון אחד מבצע מהלך בכל רגע"),
        ("דליפת קובץ הסיסמאות", "נמוכה", "גבוהה",
         "נשמרים רק hash + salt; אי אפשר לשחזר את הסיסמה המקורית"),
        ("שחקן מנסה לרמות במהלך", "בינונית", "בינונית",
         "השרת הוא הסמכות — בודק חוקיות כל מהלך לפני שמבצע אותו"),
    ])


def write_knowledge_domain(doc):
    heading(doc, "2.1 חוקי שחמט וניהול משחק", level=2)
    he_paragraph(doc,
                 "המשחק מתנהל על לוח 8×8 (64 משבצות). כל שחקן מתחיל עם 16 כלים: "
                 "8 חיילים, 2 צריחים, 2 פרשים, 2 רצים, מלכה ומלך. החוקים שמומשו "
                 "במחלקה Board:")
    for line in [
        "תנועה חוקית לפי סוג הכלי (חייל קדימה, צריח בקווים, רץ באלכסונים, פרש בצורת L, מלכה בכל הכיוונים, מלך משבצת אחת)",
        "אכילת כלי יריב על ידי מעבר למשבצת שלו",
        "קידום חייל — חייל שמגיע לשורה האחרונה, השחקן בוחר לאיזה כלי לקדם (מלכה/צריח/רץ/פרש)",
        "הכאה דרך הילוכו (en passant) — חייל יכול לאכול חייל יריב שזה עתה זינק שתי משבצות",
        "הצרחה (castling) — המלך עובר שתי משבצות לכיוון הצריח, והצריח קופץ לצד השני; הצרחה קצרה (kingside) ארוכה (queenside)",
        "שח (check) — חובה להגן על המלך; אסור לבצע מהלך שמשאיר את המלך מאוים",
        "מט (checkmate) — אין מהלך חוקי והמלך מאוים → סוף משחק, הצד התקוף מפסיד",
        "פט (stalemate) — אין מהלך חוקי אך המלך לא מאוים → תיקו",
    ]:
        he_bullet(doc, line)
    callout(doc, [
        ("סיכום חוקים שמומשו: ", "bold"),
        ("כל חוקי השחמט הבסיסיים מומשו במלואם — כולל הצרחה קצרה וארוכה, קידום חייל בבחירה "
         "והכאה דרך הילוכו. הדבר היחיד שלא מומש הוא כלל החמישים מהלכים ושלוש פעמים חזרה "
         "(תיקו לפי חזרה), שאינם קריטיים לרמת הפרויקט.", "he"),
    ])

    heading(doc, "2.2 תקשורת רשת — מודל השכבות ו-TCP", level=2)
    he_paragraph(doc,
                 "תקשורת ברשת בנויה משכבות. מודל OSI מתאר 7 שכבות. הפרויקט שלי עובד "
                 "בעיקר בשכבות הגבוהות:")
    add_table(doc, ["שכבה", "טכנולוגיה", "תפקיד בפרויקט"], [
        ("7 — אפליקציה", "פרוטוקול JSON משלי", "הודעות בין שרת ללקוח (login, move, ...)"),
        ("6 — ייצוג", "JSON (UTF-8)", "המרת ההודעות לטקסט שאפשר לשלוח (JSONUtil)"),
        ("4 — תעבורה", "TCP", "מחלקות Socket / ServerSocket; אמין ומסודר"),
        ("3 — רשת", "IP", "כתובת המחשב ברשת (למשל 127.0.0.1)"),
    ])
    he_paragraph(doc,
                 "למה TCP ולא UDP? TCP מבטיח שכל הודעה תגיע, ובסדר הנכון. בשחמט זה "
                 "קריטי — אסור שמהלך 'יאבד' או יגיע אחרי מהלך מאוחר יותר. UDP מהיר "
                 "יותר אבל לא אמין, ומתאים יותר למשחקים מהירים כמו וידאו, שבהם אובדן "
                 "הודעה בודדת לא קריטי.")
    he_paragraph(doc,
                 "כתובת IP מזהה מחשב ברשת, ו-PORT (בפרויקט: 5555) מזהה את התוכנית "
                 "הספציפית באותו מחשב. יחד הם מזהים לאן בדיוק לשלוח את ההודעה.")

    heading(doc, "2.3 גיבוב סיסמאות (Hashing)", level=2)
    he_paragraph(doc,
                 "Hash היא פונקציה חד-כיוונית: קל לחשב hash מתוך סיסמה, אבל קשה מאוד "
                 "לחזור מה-hash אל הסיסמה המקורית. בפרויקט אני משתמש ב-SHA-256, "
                 "פונקציית גיבוב סטנדרטית שמחזירה פלט באורך קבוע (256 ביט), דרך "
                 "המחלקה MessageDigest של Java.")
    he_paragraph(doc,
                 "כדי לחזק את האבטחה, מוסיפים לכל סיסמה salt — מחרוזת אקראית ייחודית "
                 "למשתמש, שנוצרת באמצעות SecureRandom. ה-salt נשמר יחד עם ה-hash. כך, "
                 "גם אם שני משתמשים בחרו את אותה סיסמה, ה-hash שלהם יהיה שונה, וזה "
                 "מקשה מאוד על תוקף שמנסה להשתמש בטבלאות hash מוכנות מראש (rainbow tables).")
    code_block(doc, [
        "// UserAuth.java - חישוב ה-hash",
        "byte[] saltBytes = new byte[16];",
        "new SecureRandom().nextBytes(saltBytes);          // salt אקראי לכל משתמש",
        "String salt = Base64.getEncoder().encodeToString(saltBytes);",
        "",
        "MessageDigest md = MessageDigest.getInstance(\"SHA-256\");",
        "byte[] digest = md.digest((salt + password).getBytes(StandardCharsets.UTF_8));",
        "// רק salt + hex(digest) נשמרים בקובץ, לא הסיסמה",
    ])
    callout(doc, [
        ("הבחנה חשובה: ", "bold"),
        ("גיבוב (hashing) הוא לא הצפנה. הצפנה היא דו-כיוונית (אפשר לפענח בחזרה), "
         "ואילו גיבוב הוא חד-כיווני. לסיסמאות משתמשים בגיבוב, כי השרת לא צריך לדעת "
         "את הסיסמה המקורית — רק לבדוק אם הסיסמה שהוזנה נותנת את אותו hash.", "he"),
    ])

    heading(doc, "2.4 ריבוי תהליכונים (Threads)", level=2)
    he_paragraph(doc, "ההבדל בין תהליך (Process) לתהליכון (Thread):")
    he_bullet(doc, "תהליך (Process): תוכנית רצה עם זיכרון נפרד משלה.")
    he_bullet(doc, "תהליכון (Thread): יחידת ריצה בתוך אותה תוכנית; כל התהליכונים חולקים את אותו זיכרון.")
    he_paragraph(doc,
                 "בשרת אני משתמש בתהליכון נפרד לכל לקוח שמתחבר. כך השרת יכול לטפל "
                 "בכמה לקוחות ובכמה משחקים בו-זמנית, בלי שלקוח אחד יחסום את האחרים.")
    he_paragraph(doc,
                 "כשכמה תהליכונים ניגשים לאותו מידע משותף (למשל אותו משחק), צריך "
                 "להגן עליו. ב-Java עשיתי זאת באמצעות המילה synchronized על מתודות "
                 "המשחק והתור — כך רק תהליכון אחד נכנס בכל פעם, והמידע לא מתקלקל.")
    code_block(doc, [
        "// ChessServer.java - תהליכון לכל לקוח",
        "Socket socket = server.accept();",
        "ClientHandler handler = new ClientHandler(socket, auth, matchmaker);",
        "new Thread(handler).start();",
    ])

    heading(doc, "2.5 שמירת נתונים בקובץ", level=2)
    he_paragraph(doc,
                 "החשבונות נשמרים בקובץ users.json בפורמט JSON. JSON הוא פורמט טקסט "
                 "קריא לשמירת מבני נתונים. בכל הרשמה חדשה הקובץ נכתב מחדש עם כל "
                 "המשתמשים, וכל משתמש שמור עם ה-salt וה-hash שלו. את הפורמט קוראת "
                 "וכותבת המחלקה UserAuth, ללא שום ספרייה חיצונית.")


def write_architecture(doc):
    heading(doc, "3.1 ארכיטקטורה כללית של המערכת", level=2)
    he_paragraph(doc, "תרשים בלוקים של רכיבי המערכת:")
    code_block(doc, [
        "+------------------+        TCP          +---------------------------+",
        "|  Client (Swing)  |<------------------->|        ChessServer        |",
        "|  ChessGUI        |   JSON over socket  |        (ServerSocket)      |",
        "+------------------+     (port 5555)     |                           |",
        "                                         |   accept loop:            |",
        "+------------------+        TCP          |   Thread per client       |",
        "|  Client (Swing)  |<------------------->|                           |",
        "|  ChessGUI        |                     |   +-------------------+   |",
        "+------------------+                     |   |  ClientHandler    |   |",
        "                                         |   +---------+---------+   |",
        "                                         |             |             |",
        "                                         |     +-------+-------+     |",
        "                                         |     |               |     |",
        "                                         |     v               v     |",
        "                                         | +--------+   +-----------+ |",
        "                                         | |Matchm. |   |   Game    | |",
        "                                         | | queue  |   | + Board   | |",
        "                                         | +--------+   | + ChessAI | |",
        "                                         |              +-----------+ |",
        "                                         |                            |",
        "                                         | +---------------------+    |",
        "                                         | |  UserAuth           |--> users.json",
        "                                         | |  (SHA-256 + salt)   |    |",
        "                                         | +---------------------+    |",
        "                                         +---------------------------+",
    ])

    heading(doc, "3.2 תרשימי זרימה", level=2)

    heading(doc, "3.2.1 הרשמה / התחברות", level=3)
    code_block(doc, [
        "Client                              Server",
        "  |                                   |",
        "  |-- {'type':'register',             |",
        "  |    'username':..., 'password':..} ->|",
        "  |                                   |  UserAuth.register()",
        "  |                                   |  salt + SHA-256 hash",
        "  |                                   |  save users.json",
        "  |<-- {'type':'auth_result',         |",
        "  |     'ok':true} --------------------|",
    ])

    heading(doc, "3.2.2 מהלך במשחק אדם מול אדם", level=3)
    code_block(doc, [
        "Client A (White)      Server                 Client B (Black)",
        "   |                    |                          |",
        "   |-- move e2->e4 ---->|                          |",
        "   |                    |  board.applyMove(...)    |",
        "   |                    |  check legal? mate?      |",
        "   |<-- state ----------|----- state ------------->|",
        "   |                    |                          |",
        "   |                    |<------ move e7->e5 ------|",
        "   |<-- state ----------|----- state ------------->|",
    ])

    heading(doc, "3.2.3 קידום חייל (בחירת כלי) — סבב שרת", level=3)
    code_block(doc, [
        "Client                 Server",
        "  |                      |",
        "  |-- move c7->c8 ------>|  board.isLegalMove()  → חוקי",
        "  |                      |  board.isPromotionMove() → כן",
        "  |<-- choose_promotion -|  (השרת מבקש לבחור כלי)",
        "  |                      |",
        "  |-- promote (choice) ->|  board.applyMove(..., choice)",
        "  |<-- state ------------|",
    ])

    heading(doc, "3.2.4 משחק נגד מחשב", level=3)
    code_block(doc, [
        "Client                 Server",
        "  |                      |",
        "  |-- play_ai (level 2) >|  new Game(human, null, new ChessAI(2))",
        "  |<-- game_start -------|",
        "  |<-- state ------------|",
        "  |                      |",
        "  |-- move e2->e4 ------>|  board.applyMove(...)",
        "  |<-- state ------------|",
        "  |                      |  ChessAI.pickMove(board)",
        "  |                      |  board.applyMove(aiMove)",
        "  |<-- state (AI move) --|",
    ])

    heading(doc, "3.3 חלוקה לחבילות וקבצים — שרת מול לקוח", level=2)
    add_table(doc, ["קובץ / חבילה", "צד", "תפקיד"], [
        ("chess.Color", "משותף", "enum של צבע השחקן (WHITE / BLACK)"),
        ("chess.pieces.*", "משותף", "מחלקת Piece המופשטת + 6 מחלקות הכלים היורשות"),
        ("chess.Board", "משותף", "לוח, חוקי תנועה, שח/מט/פט, קידום, en passant"),
        ("chess.ChessAI", "שרת", "בחירת מהלך למחשב (חמדנית)"),
        ("chess.UserAuth", "שרת", "גיבוב סיסמאות ושמירה בקובץ"),
        ("chess.JSONUtil", "משותף", "קידוד ופענוח הודעות JSON ללא ספרייה חיצונית"),
        ("chess.server.*", "שרת", "ChessServer + ClientHandler + Matchmaker + Game"),
        ("chess.client.*", "לקוח", "ChessGUI + NetworkClient (ממשק Swing)"),
        ("chess.Main", "משותף", "נקודת כניסה: הרצת שרת או לקוח"),
    ])

    heading(doc, "3.4 ניתוח אלגוריתמים מרכזיים — שקילת חלופות", level=2)

    heading(doc, "3.4.1 בדיקת חוקיות מהלך", level=3)
    he_paragraph(doc,
                 "כל כלי יודע לחשב לאן הוא 'יכול' לזוז (תנועה בסיסית). אבל מהלך חוקי "
                 "הוא מהלך שגם לא משאיר את המלך שלי בשח. כדי לבדוק זאת, מבצעים את "
                 "המהלך באופן זמני, בודקים אם המלך מאוים, ואז מחזירים את הלוח לאחור:")
    code_block(doc, [
        "// Board.java",
        "private boolean moveLeavesKingInCheck(int fc, int fr, int tc, int tr) {",
        "    Piece saved = grid[tc][tr];",
        "    grid[tc][tr] = grid[fc][fr];      // בצע זמנית",
        "    grid[fc][fr] = null;",
        "    boolean inCheck = isInCheck(turn);// בדוק שח",
        "    grid[fc][fr] = grid[tc][tr];      // החזר לאחור",
        "    grid[tc][tr] = saved;",
        "    return inCheck;",
        "}",
    ])
    he_bullet(doc, "חלופה: לחשב מראש את כל המהלכים החוקיים בכל פוזיציה ולשמור אותם — נדחתה, מסבכת את הקוד ולא נחוצה במשחק בקצב אנושי.")

    heading(doc, "3.4.2 בינה מלאכותית", level=3)
    he_paragraph(doc,
                 "ה-AI הוא 'חמדן' (greedy): לכל מהלך חוקי מחושב ניקוד לפי ערך הכלי "
                 "שנאכל (אם נאכל), בתוספת בונוס קטן על התקדמות למרכז הלוח. נבחר המהלך "
                 "עם הניקוד הגבוה ביותר (שובר שוויון אקראי).")
    he_bullet(doc, "✅ חמדנות (greedy) — נבחרה. פשוטה להבנה, מנצחת מהלכים אקראיים, ומתאימה להיקף הפרויקט.")
    he_bullet(doc, "⏳ Minimax עם Alpha-Beta — חלופה חזקה יותר שמסתכלת כמה מהלכים קדימה. נדחתה כי היא מורכבת מדי לרמת 5 יחידות ומכבידה על הקוד.")
    he_bullet(doc, "⏳ רשת נוירונים — נדחתה לחלוטין; דורשת אימון וחומרה, הרבה מעבר להיקף.")

    heading(doc, "3.4.3 פרוטוקול ההודעות", level=3)
    he_bullet(doc, "✅ JSON על TCP, שורה לכל הודעה — נבחר. קריא, קל לפענוח, וקל לבדיקה.")
    he_bullet(doc, "⏳ פרוטוקול בינארי — קומפקטי יותר, אבל קשה לקריאה ולניפוי שגיאות. נדחה.")

    heading(doc, "3.5 פרוטוקול התקשורת — מכונת מצבים", level=2)
    he_paragraph(doc,
                 "כל לקוח עובר בין מצבים: לא מחובר → מחובר (אחרי login) → ממתין בתור "
                 "או במשחק. השרת לא מאפשר לבצע פעולה שלא מתאימה למצב (למשל לשלוח "
                 "מהלך לפני שיש משחק).")


def write_implementation(doc):
    heading(doc, "4.1 תכנות מונחה עצמים — המחלקות", level=2)
    he_paragraph(doc,
                 "הפרויקט בנוי מ-19 מחלקות. הלב של התכנון מונחה-העצמים הוא מחלקת "
                 "האב המופשטת Piece ושש מחלקות הכלים שיורשות ממנה:")
    add_table(doc, ["מחלקה", "חבילה", "תפקיד"], [
        ("Piece (abstract)", "chess.pieces", "מחלקת אב מופשטת: צבע, האם זז, ומתודות getMoves/getValue/getSymbol"),
        ("Pawn / Knight / Bishop", "chess.pieces", "כלים יורשים — כל אחד מממש את getMoves שלו"),
        ("Rook / Queen / King", "chess.pieces", "כלים יורשים נוספים"),
        ("Board", "chess", "לוח 8×8, חוקי תנועה, שח/מט/פט, קידום, en passant"),
        ("ChessAI", "chess", "בחירת מהלך למחשב (חמדנית)"),
        ("UserAuth", "chess", "הרשמה/התחברות + גיבוב סיסמאות + users.json"),
        ("JSONUtil", "chess", "קידוד/פענוח JSON ללא ספרייה חיצונית"),
        ("Matchmaker", "chess.server", "תור FIFO לזיווג שחקנים"),
        ("Game", "chess.server", "משחק בודד: לוח + שני שחקנים + AI אופציונלי"),
        ("ClientHandler", "chess.server", "מטפל בלקוח אחד (מממש Runnable, רץ ב-Thread נפרד)"),
        ("ChessServer", "chess.server", "השרת הראשי: accept loop + ניהול חיבורים"),
        ("NetworkClient", "chess.client", "צד הלקוח של ה-socket + תור הודעות נכנסות"),
        ("ChessGUI", "chess.client", "ממשק Swing: התחברות, לובי, לוח"),
    ])
    he_paragraph(doc,
                 "המחלקות מדגימות עקרונות מרכזיים בתכנות מונחה עצמים: ירושה "
                 "(הכלים יורשים מ-Piece), פולימורפיזם (אותה קריאה getMoves מתנהגת "
                 "אחרת לכל כלי), הפשטה (Piece היא abstract עם מתודות מופשטות), "
                 "כימוס (כל מחלקה מסתירה את הפרטים הפנימיים שלה), והרכבה (Board "
                 "מכיל כלים, Game מכיל Board).")

    heading(doc, "4.2 דוגמה לירושה ולפולימורפיזם", level=2)
    code_block(doc, [
        "// Piece.java - מחלקת אב מופשטת",
        "public abstract class Piece {",
        "    public final Color color;",
        "    public abstract List<int[]> getMoves(Board board, int col, int row);",
        "    public abstract int getValue();",
        "}",
        "",
        "// Rook.java - יורש",
        "public class Rook extends Piece {",
        "    private static final int[][] DIRS = {{1,0},{-1,0},{0,1},{0,-1}};",
        "    public List<int[]> getMoves(Board b, int col, int row) {",
        "        return slideMoves(b, col, row, DIRS);",
        "    }",
        "}",
        "",
        "// Knight.java - יורש אחר",
        "public class Knight extends Piece {",
        "    public List<int[]> getMoves(Board b, int col, int row) {",
        "        return jumpMoves(b, col, row, OFFSETS);",
        "    }",
        "}",
    ])
    he_paragraph(doc,
                 "כאשר הלוח קורא piece.getMoves(...) הוא לא צריך לדעת איזה סוג כלי "
                 "זה — כל כלי 'יודע' בעצמו איך הוא זז. זהו פולימורפיזם.")

    heading(doc, "4.3 הצרחה — מימוש בתוך Board", level=2)
    he_paragraph(doc,
                 "ההצרחה מימושה כולה ב-Board.java ללא שינוי ב-King.java, כדי למנוע רקורסיה: "
                 "addCastlingMoves בודקת את כל חמשת תנאי ה-FIDE ומוסיפה את המהלך לרשימה "
                 "החוקית רק אם כולם מתקיימים:")
    code_block(doc, [
        "// Board.java - תנאי הצרחה",
        "private void addCastlingMoves(List<int[]> legal, int kc, int kr, Color color) {",
        "    if (grid[kc][kr].hasMoved) return;      // (1) המלך לא זז",
        "    if (isInCheck(color)) return;            // (2) המלך לא בשח",
        "    // Kingside: (3) הצריח לא זז (4) ריק בינהם (5) לא עובר תחת מתקפה",
        "    Piece kRook = grid[7][kr];",
        "    if (kRook instanceof Rook && !kRook.hasMoved",
        "            && grid[5][kr]==null && grid[6][kr]==null",
        "            && !isAttackedSimple(5,kr,opp) && !isAttackedSimple(6,kr,opp))",
        "        legal.add(new int[]{6, kr});",
        "}",
    ])

    heading(doc, "4.4 הדגשת מהלכים חוקיים (UX)", level=2)
    he_paragraph(doc,
                 "בלחיצה על כלי, הלקוח שולח get_moves לשרת. השרת מחזיר רשימת "
                 "משבצות חוקיות, והלקוח מציג:")
    he_bullet(doc, "עיגול ירוק מלא — משבצת ריקה שאפשר לזוז אליה (כולל משבצת ההצרחה)")
    he_bullet(doc, "מסגרת עגולה ירוקה — כלי יריב שאפשר לאכול")
    he_paragraph(doc, "כך השחקן רואה בדיוק אילו מהלכים חוקיים לפני שהוא לוחץ, כולל הצרחה ו-en passant.")
    code_block(doc, [
        "// ChessGUI.java - בקשת מהלכים חוקיים",
        "Map<String,Object> req = new LinkedHashMap<>();",
        "req.put(\"type\", \"get_moves\");",
        "req.put(\"col\", col);",
        "req.put(\"row\", rank);",
        "net.send(req);",
        "// התשובה מגיעה כ-{type:'moves', moves:[[6,0],[5,1],...]}",
        "// ומצוירת כנקודות ירוקות על הלוח",
    ])

    heading(doc, "4.5 קוד בטוח — טיפול בשגיאות", level=2)
    he_paragraph(doc,
                 "השרת חייב לשרוד ניתוק לקוח או הודעה לא תקינה. לכן הלולאה הראשית "
                 "של כל לקוח עטופה ב-try/catch, וכל שגיאה מטופלת בלי להפיל את השרת:")
    code_block(doc, [
        "// ClientHandler.java - run()",
        "public void run() {",
        "    try {",
        "        String line;",
        "        while ((line = in.readLine()) != null) {",
        "            handleMessage(JSONUtil.decode(line));",
        "        }",
        "    } catch (IOException ignored) {",
        "    } finally {",
        "        cleanup();   // מסיר מהתור / סוגר את ה-socket",
        "    }",
        "}",
    ])


def write_user_guide(doc):
    heading(doc, "5.1 התקנה והרצה", level=2)
    he_paragraph(doc, "דרישות: JDK 17 ומעלה מותקן. אין צורך בספריות חיצוניות.")
    code_block(doc, [
        "cd chess-bagrut-java",
        "",
        ":: טרמינל 1 - הפעלת השרת:",
        "run.bat server",
        "",
        ":: טרמינל 2 - הפעלת לקוח:",
        "run.bat client",
    ])
    he_paragraph(doc,
                 "ב-VS Code / Cursor אפשר גם להשתמש בלשונית Run and Debug ולבחור "
                 "'Chess: Run Server' ואז 'Chess: Run Client'. הקובץ run.bat מאתר "
                 "את ה-JDK אוטומטית, מקמפל את הקוד לתיקיית bin, ומריץ.")

    heading(doc, "5.2 צילומי מסך של הזרימה", level=2)
    callout(doc, [
        ("פעולה נדרשת מהתלמיד: ", "bold"),
        ("יש להריץ את המשחק ולצרף כאן צילומי מסך לפי הרשימה.", "he"),
    ])
    he_bullet(doc, "[צילום 5.2.1] מסך התחברות — לפני הזנת פרטים")
    he_bullet(doc, "[צילום 5.2.2] מסך לובי — אחרי התחברות")
    he_bullet(doc, "[צילום 5.2.3] מסך לובי — 'Looking for an opponent...'")
    he_bullet(doc, "[צילום 5.2.4] לוח המשחק — תחילת משחק עם 32 כלים ומחוון תור")
    he_bullet(doc, "[צילום 5.2.5] לוח המשחק — נקודות ירוקות מראות מהלכים חוקיים אחרי לחיצה על כלי")
    he_bullet(doc, "[צילום 5.2.6] לוח המשחק — ביצוע הצרחה (המלך ב-g1 והצריח ב-f1)")
    he_bullet(doc, "[צילום 5.2.7] לוח המשחק — דיאלוג קידום חייל (בחירת כלי)")

    heading(doc, "5.3 סוגי משתמשים", level=2)
    heading(doc, "5.3.1 שחקן אדם מול אדם", level=3)
    he_paragraph(doc,
                 "1) הרץ שרת. 2) הרץ שני לקוחות. 3) בכל לקוח הירשם/התחבר עם שם משתמש "
                 "שונה. 4) בכל לקוח לחץ 'Play vs Human'. 5) ברגע ששני השחקנים בתור, "
                 "המשחק מתחיל אוטומטית. 6) לחץ על הכלי שלך ואז על משבצת היעד.")

    heading(doc, "5.3.2 שחקן מול מחשב", level=3)
    he_paragraph(doc,
                 "במסך הלובי לחץ 'Play vs AI'. אתה משחק לבן והמחשב משחק שחור ועונה "
                 "אוטומטית אחרי כל מהלך שלך.")

    heading(doc, "5.4 הוכחת הצפנת הסיסמאות", level=2)
    he_paragraph(doc,
                 "אחרי הרשמת משתמשים, אפשר לפתוח את הקובץ users.json ולראות שאין בו "
                 "סיסמאות גלויות — רק salt ו-hash (מופרדים בנקודתיים):")
    code_block(doc, [
        "{",
        '  "alice":"9f3c1a...e7:2b8d4f...a1",',
        '  "bob":"7a1b2c...d9:55ef00...3c"',
        "}",
    ])


def write_reflection(doc):
    heading(doc, "6.1 מה למדתי על עצמי", level=2)
    callout(doc, [
        ("פעולה נדרשת מהתלמיד: ", "bold"),
        ("סעיף זה חייב להיות אישי ובמילים שלך. להלן תבנית להשראה בלבד — יש להחליף "
         "בתוכן אמיתי שלך.", "he"),
    ])
    he_paragraph(doc,
                 "[תבנית להחלפה] בפרויקט הזה גיליתי שאני נהנה/נהנית במיוחד מ___, "
                 "ושכאשר נתקלתי בקושי ב___ הצלחתי להתמודד איתו על ידי ___. "
                 "למדתי לתכנן את העבודה בשלבים, וראיתי שאני עובד/ת טוב יותר כש___.",
                 italic=True, color=MUTED)

    heading(doc, "6.2 מה למדתי מקצועית", level=2)
    he_bullet(doc, "ההבדל בין גיבוב (hash, חד-כיווני, לסיסמאות) לבין הצפנה (דו-כיוונית).")
    he_bullet(doc, "למה salt חשוב — בלעדיו, סיסמאות זהות מקבלות hash זהה.")
    he_bullet(doc, "איך עובדים סוקטים ב-TCP ב-Java: ServerSocket, accept, Socket, connect.")
    he_bullet(doc, "למה צריך תהליכון לכל לקוח, ולמה צריך synchronized על מידע משותף.")
    he_bullet(doc, "ירושה, הפשטה (abstract) ופולימורפיזם — מחלקת Piece אחת חסכה לי המון קוד כפול.")
    he_bullet(doc, "החשיבות של טיפול בשגיאות (try/catch) כדי שהשרת לא יקרוס.")
    he_bullet(doc, "הצרחה — למדתי שחוקי FIDE כוללים 5 תנאים מצטברים: המלך והצריח לא זזו, הנתיב ריק, המלך לא בשח, ולא עובר דרך שדה מאוים.")
    he_bullet(doc, "UX חשוב לא פחות מהלוגיקה — הוספת נקודות ירוקות (הדגשת מהלכים) חיסלה שאלות בנוסח 'למה אי אפשר לזוז?'.")

    heading(doc, "6.3 מה הייתי משנה אם הייתי מתחיל מחדש", level=2)
    he_bullet(doc, "להוסיף כלל חמישים מהלכים ותיקו לפי שלוש חזרות, להשלמת חוקי FIDE.")
    he_bullet(doc, "לשפר את ה-AI כך שיסתכל שני מהלכים קדימה (Minimax).")
    he_bullet(doc, "להוסיף שמירת היסטוריית משחקים וטבלת ניקוד.")
    he_bullet(doc, "להוסיף שעון לכל שחקן (משחק מהיר).")


def write_bibliography(doc):
    he_paragraph(doc,
                 "להלן המקורות שבהם נעזרתי, ולכל מקור — תרומתו לפרויקט.")
    sources = [
        ("[1] תיעוד java.net.Socket / ServerSocket",
         "https://docs.oracle.com/en/java/javase/17/docs/api/java.base/java/net/Socket.html",
         "עזר להבין איך פותחים שרת (ServerSocket + accept) ואיך לקוח מתחבר (Socket + connect), "
         "ומה ההבדל בין TCP ל-UDP."),
        ("[2] תיעוד Java Concurrency (Thread, synchronized)",
         "https://docs.oracle.com/javase/tutorial/essential/concurrency/",
         "עזר להבין Thread ו-synchronized, ואיך מריצים תהליכון נפרד לכל לקוח."),
        ("[3] תיעוד java.security.MessageDigest",
         "https://docs.oracle.com/en/java/javase/17/docs/api/java.base/java/security/MessageDigest.html",
         "עזר להבין איך מחשבים SHA-256 ב-Java ואיך מוסיפים salt עם SecureRandom."),
        ("[4] תיעוד Java Swing (JFrame, JPanel, paintComponent)",
         "https://docs.oracle.com/javase/tutorial/uiswing/",
         "ספריית הממשק הגרפי. למדתי על ציור מותאם ב-paintComponent ועל Timer לעדכון תקופתי."),
        ("[5] חוקי השחמט הרשמיים (FIDE)",
         "https://www.fide.com/FIDE/handbook/LawsOfChess.pdf",
         "מקור החוקים שמימשתי: תנועת כל כלי, שח, מט, פט, קידום חייל והכאה דרך הילוכו."),
        ("[6] מדריך JSON",
         "https://www.json.org/json-en.html",
         "פורמט ההודעות בין השרת ללקוח. עזר להבין את מבנה ה-JSON שמימשתי ב-JSONUtil."),
    ]
    for title, url, desc in sources:
        heading(doc, title, level=2)
        he_paragraph(doc, url, color=MUTED, italic=True)
        he_paragraph(doc, desc)


def write_source_appendix(doc):
    he_paragraph(doc,
                 "תדפיס מסודר של כל קבצי הקוד בפרויקט (Java SE).")
    for relative in SOURCE_FILES:
        path = REPO_ROOT / relative
        if not path.exists():
            continue
        heading(doc, relative, level=2)
        lines = path.read_text(encoding="utf-8").splitlines()
        if not lines:
            lines = ["// (קובץ ריק)"]
        code_block(doc, lines)


def write_protocol_appendix(doc):
    heading(doc, "ב.1 הודעות לקוח → שרת", level=2)
    add_table(doc, ["סוג הודעה", "תיאור", "שדות"], [
        ("register", "הרשמת משתמש חדש", "username, password"),
        ("login", "התחברות לחשבון קיים", "username, password"),
        ("play_human", "הצטרפות לתור משחק מול אדם", "(אין)"),
        ("play_ai", "התחלת משחק נגד מחשב", "level"),
        ("cancel_wait", "ביטול ההמתנה בתור", "(אין)"),
        ("move", "ביצוע מהלך", "from_col, from_row, to_col, to_row"),
        ("promote", "בחירת כלי לקידום חייל", "choice (Q/R/B/N)"),
        ("resign", "כניעה", "(אין)"),
    ])

    heading(doc, "ב.2 הודעות שרת → לקוח", level=2)
    add_table(doc, ["סוג הודעה", "תיאור", "שדות"], [
        ("auth_result", "תוצאת הרשמה/התחברות", "ok, message, username"),
        ("info", "הודעת מידע (למשל 'ממתין ליריב')", "message"),
        ("error", "הודעת שגיאה (למשל מהלך לא חוקי)", "message"),
        ("game_start", "המשחק התחיל", "color, opponent"),
        ("state", "עדכון מצב הלוח", "board, turn, status, winner, last_move"),
        ("choose_promotion", "בקשה לבחור כלי לקידום", "(אין)"),
    ])

    heading(doc, "ב.3 דוגמת זרימה", level=2)
    code_block(doc, [
        'C -> S  {"type":"register","username":"alice","password":"1234"}',
        'S -> C  {"type":"auth_result","ok":true,"username":"alice"}',
        "",
        'C -> S  {"type":"play_ai","level":2}',
        'S -> C  {"type":"game_start","color":"white","opponent":"AI"}',
        'S -> C  {"type":"state","board":[...],"turn":"white",...}',
        "",
        'C -> S  {"type":"move","from_col":4,"from_row":1,"to_col":4,"to_row":3}',
        'S -> C  {"type":"state","board":[...],"turn":"black","last_move":[4,1,4,3]}',
        'S -> C  {"type":"state","board":[...],"turn":"white",...}   // מהלך המחשב',
    ])


def write_qa_appendix(doc):
    he_paragraph(doc,
                 "להלן שאלות תיאורטיות אפשריות מהבוחן עם תשובות מוכנות, לפי נושא.")

    heading(doc, "ג.1 תקשורת", level=2)
    qa = [
        ("מה ההבדל בין TCP ל-UDP?",
         "TCP אמין ומסודר (כל הודעה מגיעה ובסדר הנכון) אבל איטי יותר; UDP מהיר אבל "
         "לא אמין. בחרתי TCP כי בשחמט אסור שמהלך יאבד."),
        ("מה זה כתובת IP ומה זה PORT?",
         "כתובת IP מזהה מחשב ברשת. PORT מזהה תוכנית מסוימת באותו מחשב (אצלי 5555)."),
        ("מהם השלבים בפתיחת שרת ב-Java?",
         "יוצרים ServerSocket עם פורט, ואז בלולאה קוראים accept() שמחזיר Socket לכל "
         "לקוח שמתחבר."),
        ("מה עושה הלקוח כדי להתחבר?",
         "יוצר new Socket(host, port) — זה מבצע connect לשרת לפי כתובת ה-IP והפורט."),
    ]
    for q, a in qa:
        he_segments(doc, [("ש: ", "bold"), (q, "he")], size=12)
        he_segments(doc, [("ת: ", "accent"), (a, "he")], size=12, space_after=8)

    heading(doc, "ג.2 גיבוב והצפנה", level=2)
    qa = [
        ("מה זה hash?",
         "פונקציה חד-כיוונית שמקבלת קלט ומחזירה פלט בגודל קבוע. קל לחשב, קשה מאוד "
         "להפוך בחזרה. דוגמה: SHA-256."),
        ("מה ההבדל בין גיבוב להצפנה?",
         "הצפנה דו-כיוונית — אפשר לפענח בחזרה עם מפתח. גיבוב חד-כיווני — אי אפשר "
         "לשחזר את הקלט. לסיסמאות משתמשים בגיבוב."),
        ("מה זה salt ולמה הוא חשוב?",
         "מחרוזת אקראית שמתווספת לסיסמה לפני הגיבוב. בלעדיו, שתי סיסמאות זהות "
         "מקבלות hash זהה, וזה מאפשר התקפה בעזרת טבלאות מוכנות (rainbow tables)."),
        ("איך נשמרת הסיסמה בפרויקט שלך?",
         "לא נשמרת הסיסמה עצמה. נשמרים רק ה-salt וה-hash של (salt + סיסמה) בקובץ "
         "users.json. בהתחברות מחשבים שוב את ה-hash ומשווים."),
    ]
    for q, a in qa:
        he_segments(doc, [("ש: ", "bold"), (q, "he")], size=12)
        he_segments(doc, [("ת: ", "accent"), (a, "he")], size=12, space_after=8)

    heading(doc, "ג.3 מערכות הפעלה ותהליכונים", level=2)
    qa = [
        ("מה ההבדל בין Process ל-Thread?",
         "Process הוא תוכנית עם זיכרון נפרד. Thread הוא יחידת ריצה בתוך תוכנית, "
         "וכל ה-threads חולקים את אותו זיכרון."),
        ("למה השתמשת ב-Thread לכל לקוח?",
         "כדי שהשרת יוכל לטפל בכמה לקוחות במקביל, בלי שלקוח אחד יחסום את האחרים."),
        ("מה זה synchronized ולמה צריך אותו?",
         "מילת מפתח שמבטיחה שרק תהליכון אחד יריץ את המתודה בו-זמנית, כדי שמידע "
         "משותף (כמו מצב המשחק או התור) לא יתקלקל. זה התחליף ל-Lock."),
    ]
    for q, a in qa:
        he_segments(doc, [("ש: ", "bold"), (q, "he")], size=12)
        he_segments(doc, [("ת: ", "accent"), (a, "he")], size=12, space_after=8)

    heading(doc, "ג.4 תכנות מונחה עצמים", level=2)
    qa = [
        ("מה זה ירושה?",
         "מנגנון שבו מחלקה (יורשת) מקבלת תכונות ופעולות ממחלקת אב. אצלי: Pawn, "
         "Rook וכו' יורשים מ-Piece בעזרת extends."),
        ("מה זה פולימורפיזם?",
         "אותה פעולה (getMoves) מתנהגת אחרת בכל מחלקה יורשת. הלוח קורא לה בלי "
         "לדעת איזה כלי זה בדיוק."),
        ("מה זה מחלקה מופשטת (abstract)?",
         "מחלקה שאי אפשר ליצור ממנה עצם ישירות, ויכולה להגדיר מתודות מופשטות "
         "שהיורשים חייבים לממש. אצלי Piece היא abstract עם getMoves מופשטת."),
        ("מה זה כימוס (Encapsulation)?",
         "הסתרת הפרטים הפנימיים של מחלקה וחשיפת ממשק נקי. למשל Board חושף applyMove "
         "ומסתיר איך בדיוק נשמר הלוח (grid פרטי)."),
    ]
    for q, a in qa:
        he_segments(doc, [("ש: ", "bold"), (q, "he")], size=12)
        he_segments(doc, [("ת: ", "accent"), (a, "he")], size=12, space_after=8)


if __name__ == "__main__":
    build_document()
