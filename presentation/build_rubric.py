from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_helpers import (
    ACCENT, ACCENT_DARK, DARK, MUTED,
    add_page_numbers, add_table, he_bullet, he_paragraph, heading,
    new_document, page_break, set_update_fields_on_open, warn_box,
    word_postprocess_rtl, _set_bidi,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = REPO_ROOT / "מחוון-תיק-פרויקט.docx"


def write_toc(doc):
    heading(doc, "תוכן עניינים", level=1)
    he_paragraph(doc,
                 "להלן תוכן העניינים האוטומטי של Word. כדי לעדכן אותו לאחר שינויים: "
                 "לחץ קליק־ימני על הטבלה ובחר 'Update Field' (או F9).",
                 italic=True, color=MUTED)

    p = doc.add_paragraph()
    _set_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\p " "'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = p.add_run()._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

    he_paragraph(doc, "תוכן עניינים — לחץ F9 ב-Word כדי לעדכן", italic=True, color=MUTED)


def build_document() -> None:
    doc = new_document()
    add_page_numbers(doc)

    write_cover(doc)
    page_break(doc)

    write_toc(doc)
    page_break(doc)

    write_part_zero(doc)
    page_break(doc)

    write_part_a(doc)
    page_break(doc)

    write_part_b(doc)
    page_break(doc)

    write_score_summary(doc)
    page_break(doc)

    write_oral_rubric(doc)
    page_break(doc)

    write_examiner_notes(doc)

    set_update_fields_on_open(doc)
    doc.save(str(OUT_DOCX))
    word_postprocess_rtl(OUT_DOCX)
    print(f"wrote rubric docx ({OUT_DOCX.stat().st_size:,} bytes)")


def write_cover(doc):
    he_paragraph(doc, "", space_before=60)
    he_paragraph(doc, "מחוון לבדיקת תיק הפרויקט", size=24, bold=True,
                 color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    he_paragraph(doc, "משחק שחמט מאובטח בין שני משתתפים (Java)", size=18, bold=True,
                 color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    he_paragraph(doc, "פרויקט גמר במדעי המחשב — 5 יחידות לימוד", size=12, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    heading(doc, "פרטי התלמיד", level=2)
    add_table(doc, ["שדה", "ערך"], [
        ("שם בית הספר", "_____________________"),
        ("סמל מוסד", "_____________________"),
        ("שם התלמיד", "_____________________"),
        ('מספר ת"ז', "_____________________"),
        ("נושא העבודה", "משחק שחמט מאובטח בין שני משתתפים — Java (שרת/לקוח + הצפנת סיסמאות)"),
        ("תאריך הערכה", "_____________________"),
        ("שם הבוחן", "_____________________"),
    ], col_widths=[5.5, 11.0])


def write_part_zero(doc):
    heading(doc, "חלק 0 — דרישות חובה בפרויקט (תנאי סף)", level=1)
    warn_box(doc,
             "הערה: המחוון בנוי לפי ששת הדרישות שהוגדרו לפרויקט. דרישות 1, 2, 5, 6 "
             "הן חובה; דרישות הבונוס (ריבוי לקוחות במקביל ובינה מלאכותית) מומשו גם הן.")

    add_table(doc, ["#", "דרישה", "סטטוס", "מיקום בפרויקט / הוכחה"], [
        ("1", "מימוש לפחות 2 מחלקות שונות הכוללות עצמים ופעולות",
         "כן (חורג)",
         "19 מחלקות. ירושה והפשטה: Piece (abstract) + Pawn/Knight/Bishop/Rook/Queen/King. "
         "בנוסף Board, ChessAI, UserAuth, JSONUtil, Matchmaker, Game, ChessServer, "
         "ClientHandler, NetworkClient, ChessGUI (תיק §4.1)"),

        ("2", "מערכת שרת/לקוח + תור לקוחות (ריבוי לקוחות במקביל = בונוס)",
         "כן + בונוס",
         "ChessServer (ServerSocket+accept) + NetworkClient (Socket connect). תור FIFO "
         "במחלקת Matchmaker. נימה נפרדת לכל לקוח (Thread) = ריבוי לקוחות במקביל"),

        ("3", "שימוש בבינה מלאכותית (בונוס)",
         "כן בונוס",
         "מחלקת ChessAI: בחירת מהלך חמדנית עם הערכת חומר ובונוס מרכז"),

        ("4", "מערכת הפעלה — ללא שינוי",
         "כן",
         "Java SE בלבד (JDK 17), ללא ספריות חיצוניות; רץ על Windows/Linux/macOS"),

        ("5", "הצפנת סיסמאות בלבד (ללא הצפנת מידע רגיש בתקשורת)",
         "כן",
         "UserAuth: SHA-256 (MessageDigest) + salt אקראי לכל משתמש (SecureRandom). "
         "הסיסמה לא נשמרת — רק salt + hash ב-users.json. התקשורת JSON רגיל"),

        ("6", "ממשק משתמש — ללא שינוי",
         "כן",
         "ממשק גרפי Swing ב-ChessGUI: שלושה מסכים (התחברות, לובי, לוח משחק לחיץ)"),
    ])

    he_paragraph(doc, "סיכום תנאי סף: כל הדרישות מולאו, כולל שני מרכיבי הבונוס.",
                 bold=True, color=ACCENT, size=13)


def write_part_a(doc):
    heading(doc, "חלק א' — מראה התיק (15%)", level=1)
    add_table(doc, ["מרכיב", "תיאור הדרישות", "משקל", "ניקוד", "סטטוס", "הערות"], [
        ("שער פתיחה", "על פי התבנית", "2%", "___", "כן",
         "שער מולא: שם פרויקט, תלמיד, ת\"ז, סמל מוסד, ביה\"ס, מורה, תאריך"),
        ("תוכן עניינים", "מקושר לפרקים", "2%", "___", "כן",
         "Word יוצר תוכן עניינים אוטומטי המקושר ל-Heading 1/2/3"),
        ("גופן אחיד", "לכלל התיק", "2%", "___", "כן",
         "Arial 12 לטקסט, Consolas 9 לקטעי קוד"),
        ("כותרות", "היררכיה ועיצוב עקבי", "3%", "___", "כן",
         "Heading 1/2/3 בצבעים עקביים"),
        ("מספרי עמוד", "בכל העמודים", "3%", "___", "כן",
         "מספור עמודים בכותרת התחתונה"),
        ("עימוד אחיד", "שוליים ומרווחים", "3%", "___", "כן",
         "שוליים 2 ס\"מ, רווח שורות 1, יישור לימין (RTL)"),
    ])
    he_paragraph(doc, "סה\"כ חלק א': ___ / 15%", bold=True, color=ACCENT)


def write_part_b(doc):
    heading(doc, "חלק ב' — תוכן התיק (85%)", level=1)

    heading(doc, "1. מבוא (ייזום, אפיון) — 13%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("תיאור הרעיון והמוטיבציה", "כן", "תיק §1.1"),
        ("ייזום: זיהוי הצורך", "כן", "תיק §1.2"),
        ("אפיון פונקציונלי", "כן", "תיק §1.3 — 10 דרישות"),
        ("אפיון לא-פונקציונלי", "כן", "תיק §1.4 — 5 דרישות איכות"),
        ("לוח זמנים הגיוני", "כן", "תיק §1.5 — אבני דרך"),
        ("ניהול סיכונים", "כן", "תיק §1.6 — 4 סיכונים עיקריים"),
    ])
    he_paragraph(doc, "ניקוד מבוא: ___ / 13%", bold=True, color=ACCENT)

    heading(doc, "2. תיאור תחום הידע — פרק מילולי (ניתוח) — 10%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("חוקי שחמט וניהול משחק", "כן", "תיק §2.1 — כולל קידום, en passant והצרחה"),
        ("תקשורת רשת — שכבות ו-TCP", "כן", "תיק §2.2"),
        ("ערבול סיסמאות (SHA-256 + salt)", "כן", "תיק §2.3"),
        ("ריבוי נימות", "כן", "תיק §2.4 — Thread + synchronized"),
        ("שמירת נתונים בקובץ", "כן", "תיק §2.5"),
    ])
    he_paragraph(doc, "ניקוד ניתוח: ___ / 10%", bold=True, color=ACCENT)

    heading(doc, "3. מבנה / ארכיטקטורה (העיצוב) — 25%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("שרטוט ארכיטקטורה", "כן", "תיק §3.1 — תרשים בלוקים"),
        ("תרשימי זרימת מידע", "כן", "תיק §3.2 — 4 דיאגרמות (כולל קידום, הצרחה, הדגשת מהלכים)"),
        ("הבחנה בין מודולי שרת ולקוח", "כן", "תיק §3.3 — טבלת חבילות"),
        ("ניתוח אלגוריתמים + שקילת חלופות", "כן", "תיק §3.4 — חוקיות מהלך, AI, פרוטוקול"),
        ("פרוטוקול תקשורת הגיוני", "כן", "תיק §3.5 — מכונת מצבים"),
    ])
    he_paragraph(doc, "ניקוד ארכיטקטורה: ___ / 25%", bold=True, color=ACCENT)

    heading(doc, "4. מימוש הפרויקט (הקוד) — 31%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("תכנות מונחה עצמים — מחלקות שהתלמיד יצר", "כן",
         "19 מחלקות, כולל ירושה, הפשטה ופולימורפיזם (תיק §4.1-4.2)"),
        ("חלוקה הגיונית לחבילות/קבצים", "כן",
         "chess / chess.pieces / chess.server / chess.client"),
        ("קוד כתוב היטב", "כן",
         "שמות משמעותיים, מתודות קצרות"),
        ("חלוקה לפעולות", "כן",
         "כל מחלקה עם פעולות ברורות וממוקדות"),
        ("חלוקה ברורה בין קוד שרת ללקוח", "כן",
         "חבילת server מול client; pieces/Board משותפים"),
        ("חוקי FIDE מתקדמים — הצרחה", "כן",
         "Board.addCastlingMoves: 5 תנאים (לא זזו, נתיב ריק, לא שח, לא עובר מתקפה) — תיק §4.3"),
        ("UX — הדגשת מהלכים חוקיים", "כן",
         "get_moves/moves round-trip; נקודות ירוקות על הלוח בלחיצת כלי — תיק §4.4"),
        ("קוד בטוח — try/catch, שרת יציב", "כן",
         "ClientHandler.run עטוף ב-try/catch/finally; ניתוק לקוח מטופל (תיק §4.5)"),
    ])
    he_paragraph(doc, "ניקוד מימוש: ___ / 31%", bold=True, color=ACCENT)

    heading(doc, "5. מדריך למשתמש — 10%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("הסבר התקנה והרצה", "כן", "תיק §5.1 — JDK 17 + run.bat"),
        ("צילומי מסך של הזרימה", "לפעולה", "תיק §5.2 — מקום מסומן ל-7 צילומים שעל התלמיד להוסיף (כולל הצרחה ונקודות ירוקות)"),
        ("הסברים לכל סוגי המשתמשים", "כן", "תיק §5.3 — אדם מול אדם, מול מחשב"),
        ("הוכחת הצפנת הסיסמאות", "כן", "תיק §5.4 — תוכן users.json"),
    ])
    he_paragraph(doc, "ניקוד מדריך: ___ / 10%", bold=True, color=ACCENT)

    heading(doc, "6. סיכום אישי / רפלקציה — 6%", level=2)
    warn_box(doc, 'אסור להסתפק ב"תודה ונהניתי". יש לפרט מה למדת על עצמך ובכלל.')
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("מה למדתי על עצמי", "לפעולה", "תיק §6.1 — תבנית; חייב מילוי אישי ע\"י התלמיד"),
        ("מה למדתי מקצועית", "כן", "תיק §6.2 — לקחים מקצועיים"),
        ("מה הייתי משנה", "כן", "תיק §6.3 — שיפורים עתידיים"),
    ])
    he_paragraph(doc, "ניקוד רפלקציה: ___ / 6%", bold=True, color=ACCENT)

    heading(doc, "7. ביבליוגרפיה — 5%", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("סקר ספרות (לא רק רשימת קישורים)", "כן",
         "תיק §7 — 6 מקורות עם הסבר תרומה לפרויקט"),
    ])
    he_paragraph(doc, "ניקוד ביבליוגרפיה: ___ / 5%", bold=True, color=ACCENT)

    heading(doc, "8. נספחים", level=2)
    add_table(doc, ["תת-סעיף", "סטטוס", "מיקום / הערות"], [
        ("תדפיס הקוד", "כן", "נספח א' — תדפיס כל קבצי הקוד (19 קבצים)"),
        ("טבלת הודעות הפרוטוקול", "כן", "נספח ב' — הודעות + דוגמת זרימה"),
        ("שאלות תיאורטיות לבחינה", "כן", "נספח ג' — תקשורת, ערבול, OS, OOP"),
    ])

    heading(doc, "בונוס — עד 10%", level=2)
    add_table(doc, ["מרכיב בונוס", "סטטוס בפרויקט", "ניקוד מומלץ"], [
        ("ריבוי לקוחות במקביל", "כן נימה לכל לקוח (Thread)", "___"),
        ("בינה מלאכותית", "כן ChessAI חמדנית", "___"),
        ("היקף וקריאות הקוד", "כן מבנה מסודר ב-OOP עם חבילות; הצרחה + הדגשת מהלכים", "___"),
    ])
    he_paragraph(doc, "ניקוד בונוס: ___ / 10%", bold=True, color=ACCENT)


def write_score_summary(doc):
    heading(doc, "חישוב ציון סופי", level=1)
    add_table(doc, ["מרכיב", "משקל", "ניקוד"], [
        ("חלק א' (מראה התיק)", "15%", "___"),
        ("חלק ב' (תוכן התיק):", "85%", "___"),
        ("  ↳ מבוא", "13%", "___"),
        ("  ↳ ניתוח", "10%", "___"),
        ("  ↳ ארכיטקטורה", "25%", "___"),
        ("  ↳ מימוש", "31%", "___"),
        ("  ↳ מדריך משתמש", "10%", "___"),
        ("  ↳ רפלקציה", "6%", "___"),
        ("  ↳ ביבליוגרפיה", "5%", "___"),
        ("סה\"כ לפני בונוס", "100%", "___"),
        ("בונוס", "+10%", "___"),
        ("ציון סופי (כולל בונוס, חסום ב-100)", "—", "___"),
    ])


def write_oral_rubric(doc):
    heading(doc, "מחוון לבדיקת הפרויקט בזמן הבחינה", level=1)
    warn_box(doc, 'הפרויקט חייב לעבוד במהלך הבדיקה. לא סרטון, לא מצגת, ולא "אתמול הכל עבד".')

    add_table(doc, ["מרכיב", "משקל", "ניקוד", "קריטריונים", "הפניות בפרויקט"], [
        ("הצגה ושליטה בפרויקט", "15%", "___",
         "מציג כל חלק, מפרט, מריץ, שולט בתהליך",
         "שרת + שני לקוחות במקביל / לקוח מול מחשב"),
        ("פרויקט עובד מקצה לקצה", "30%", "___",
         "העברת מידע מקצה לקצה ביכולת אחת לפחות",
         "register → login → play_human → game_start → move → מט"),
        ("שליטה בקוד", "30%", "___",
         "ניווט, הסבר, איך נשמרים ונשלפים נתונים",
         "users.json + UserAuth (load/save) + ערבול SHA-256"),
        ("שליטה בחומר התיאורטי", "20%", "___",
         "תקשורת, ערבול, מערכת הפעלה, OOP",
         "ראה שאלות לדוגמה בנספח ג' בתיק"),
        ("בונוס/סייבר", "5%", "___",
         "הצפנת סיסמאות, יציבות השרת",
         "SHA-256 + salt; try/catch בשרת"),
    ])
    he_paragraph(doc, "סה\"כ ציון בחינה: ___ / 100%", bold=True, color=ACCENT)


def write_examiner_notes(doc):
    heading(doc, "הערות לבוחן — איך להריץ את הפרויקט", level=1)
    he_bullet(doc, "קוד מקור: תיקיית chess-bagrut-java (Java SE, 19 מחלקות תחת src/chess)")
    he_bullet(doc, "אין תלויות חיצוניות — JDK 17 ומעלה בלבד")
    he_bullet(doc, "הרצת שרת: run.bat server")
    he_bullet(doc, "הרצת לקוח: run.bat client")
    he_bullet(doc, "הוכחת ריבוי לקוחות (בונוס): שרת + שני לקוחות מקבילים שמשחקים זה נגד זה")
    he_bullet(doc, "הוכחת בינה מלאכותית (בונוס): לקוח אחד בוחר 'Play vs AI'")
    he_bullet(doc, "הוכחת הצפנת סיסמאות: התבוננות ב-users.json אחרי הרשמה — אין סיסמה גלויה")

    heading(doc, "סטטוס ההפקה", level=1)
    he_paragraph(doc,
                 "הפקה אוטומטית מחדש: הרץ python presentation/build_portfolio.py "
                 "ו-python presentation/build_rubric.py מתוך תיקיית chess-bagrut-java — "
                 "זה ייצר מחדש את שני קבצי ה-docx על סמך הקוד הנוכחי.")

    heading(doc, "מה כבר נכלל בתיק הפורמלי (תיק-פרויקט.docx)", level=2)
    he_bullet(doc, "שער פתיחה עם כל השדות + תוכן עניינים מקושר (F9 לעדכון)")
    he_bullet(doc, "פרק 1 — מבוא: רעיון, ייזום, 10 דרישות פונקציונליות, 5 לא-פונקציונליות, לו\"ז, סיכונים")
    he_bullet(doc, "פרק 2 — ניתוח: חוקי שחמט, תקשורת ו-TCP, ערבול סיסמאות, נימות, קבצים")
    he_bullet(doc, "פרק 3 — ארכיטקטורה: תרשים בלוקים, 4 תרשימי זרימה, חלוקת חבילות, ניתוח אלגוריתמים")
    he_bullet(doc, "פרק 4 — מימוש: 19 מחלקות, ירושה/הפשטה/פולימורפיזם, הצרחה (Board.addCastlingMoves), הדגשת מהלכים (get_moves/moves), קוד בטוח")
    he_bullet(doc, "פרק 5 — מדריך משתמש: התקנה, הרצה, סוגי משתמשים, הוכחת הצפנה")
    he_bullet(doc, "פרק 6 — סיכום אישי + פרק 7 — ביבליוגרפיה")
    he_bullet(doc, "נספחים א'-ג': תדפיס קוד, טבלת פרוטוקול, שאלות לבחינה")

    heading(doc, "מה נשאר על התלמיד (ידני)", level=2)
    warn_box(doc,
             "§5.2 — צילומי מסך: יש להריץ את הממשק ולצלם 7 מסכים "
             "(התחברות, לובי, המתנה, תחילת משחק, נקודות ירוקות של מהלכים, ביצוע הצרחה, דיאלוג קידום חייל).")
    warn_box(doc,
             "§6.1 — מה למדתי על עצמי: סעיף אישי שחייב להיות במילים שלך. "
             "יש תבנית בתיק — להחליף ב-2-3 פסקאות אמיתיות.")
    warn_box(doc,
             'מילוי השדות בשער: סמל מוסד, ת"ז, ביה"ס, מורה, תאריך.')


if __name__ == "__main__":
    build_document()
